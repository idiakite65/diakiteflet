from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file, jsonify
import psycopg2
from database_manager import DB_CONFIG
import hashlib
from datetime import timedelta, datetime
import subprocess
import threading
import os
import webbrowser
import time
import logging
import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml
#from docx.oxml.shared import nsdecls
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from datetime import datetime
import pandas as pd
from io import BytesIO
from docxtpl import DocxTemplate
from tempfile import NamedTemporaryFile
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
from pathlib import Path
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import zipfile
import traceback
from werkzeug.utils import secure_filename
# احصل على المسار الأساسي للمشروع
BASE_DIR = Path(__file__).parent

app = Flask(__name__)
app.secret_key = 'your_very_strong_secret_key_here'
#app.secret_key = '0f018f96f5582fea47a52442f7eaca587db3b2165a6b0c1e410347d69e0fa147'
app.config['SESSION_PERMANENT'] = True
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(minutes=30)
#csrf = CSRFProtect(app)


def get_db_connection():
    conn = psycopg2.connect(**DB_CONFIG)
    return conn

def open_browser():
    """فتح المتصفح تلقائيًا بعد بدء الخادم"""
    time.sleep(1)  # انتظر ثانية واحدة لضمان بدء الخادم
    webbrowser.open_new('http://127.0.0.1:5000/')

@app.route('/', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        try:
            conn = get_db_connection()
            cur = conn.cursor()         
            cur.execute(
                "SELECT * FROM utilisateurs WHERE login = %s AND password = md5(%s)",
                (username, password)
            )
            user = cur.fetchone()
            
            if user:
                session['logged_in'] = True
                session['username'] = username
                flash('مرحباً بك {}'.format(username), 'success')
                return redirect(url_for('annees'))
            else:
                flash('اسم المستخدم أو كلمة المرور غير صحيحة', 'danger')
                
        except Exception as e:
            flash('حدث خطأ في النظام: {}'.format(str(e)), 'danger')
        finally:
            if conn:
                conn.close()
    
    return render_template('login.html')
#----------------------------ANNEES-------------
@app.route('/annees', methods=['GET', 'POST'])
def annees():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    conn = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()        
        # جلب البيانات للقوائم المنسدلة
        cur.execute('SELECT id_ann, annee_sc, nom_cnt FROM ann_scolaire natural join nom_centere ORDER BY id_ann DESC')
        annees_options = [f"{row[0]} - {row[1]}" for row in cur.fetchall()]      
        # جلب بيانات الأعوام الدراسية للعرض
        cur.execute('SELECT id_ann, num_ann, annee_sc FROM ann_scolaire ORDER BY id_ann DESC')
        annees_data = cur.fetchall()
        
        if request.method == 'POST':
            # معالجة إضافة/تحديث السنة الدراسية
            action = request.form.get('action')
            
            if action == 'add':
                num_ann = request.form.get('numann')
                annee_sc = request.form.get('anneesc')
                
                if not num_ann or not annee_sc:
                    flash("الرجاء ملء جميع الحقول المطلوبة", "warning")
                else:
                    cur.execute("INSERT INTO ann_scolaire(num_ann, annee_sc, id_cnt) VALUES (%s, %s, %s)", 
                               (num_ann, annee_sc, 1))
                    conn.commit()
                    flash("تمت إضافة السنة الدراسية بنجاح", "success")
                return redirect(url_for('annees'))
            
            elif action == 'update':
                id_ann = request.form.get('id_ann')
                num_ann = request.form.get('numann')
                annee_sc = request.form.get('anneesc')
                
                if not id_ann or not num_ann or not annee_sc:
                    flash("الرجاء ملء جميع الحقول المطلوبة", "warning")
                else:
                    cur.execute("UPDATE ann_scolaire SET num_ann=%s, annee_sc=%s WHERE id_ann=%s", 
                               (num_ann, annee_sc, id_ann))
                    conn.commit()
                    flash("تم تحديث السنة الدراسية بنجاح", "success")
                return redirect(url_for('annees'))
    
    except Exception as e:
        if conn:
            conn.rollback()
        flash(f"حدث خطأ: {str(e)}", "danger")
        app.logger.error(f'Error in annees: {str(e)}')
    finally:
        if conn:
            conn.close()
    
    return render_template('annees.html', 
                         annees_options=annees_options, 
                         annees_data=annees_data,
                         selected_annee=session.get('annee_info', {}).get('annee_sc', ''))

@app.route('/select_annee', methods=['POST'])
def select_annee():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    selected_annee = request.form.get('selected_annee')
    if not selected_annee:
        flash("الرجاء اختيار سنة دراسية", "warning")
        return redirect(url_for('annees'))
    
    try:
        annee_id, annee_sc = selected_annee.split(' - ', 1)
        
        conn = get_db_connection()
        with conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT id_ann, annee_sc, nom_cnt 
                    FROM ann_scolaire NATURAL JOIN nom_centere 
                    WHERE id_ann = %s
                """, (annee_id,))
                row = cur.fetchone()
                
                if row:
                    session['annee_info'] = {
                        'id_ann': row[0],
                        'annee_sc': row[1],
                        'nom_cnt': row[2]
                    }
                    flash(f"تم اختيار السنة الدراسية: {row[1]}", "success")
                    return redirect(url_for('menu'))
    
    except Exception as e:
        flash(f"حدث خطأ في النظام: {str(e)}", "danger")
        return redirect(url_for('annees'))

@app.route('/delete_annee/<int:id>')
def delete_annee(id):
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    try:
        conn = get_db_connection()
        with conn:
            with conn.cursor() as cur:               
                # حذف السنة الدراسية
                cur.execute("DELETE FROM ann_scolaire WHERE id_ann=%s", (id,))
                conn.commit()
                flash("تم حذف السنة الدراسية بنجاح", "success")
    
    except Exception as e:
        flash(f"حدث خطأ أثناء حذف السنة الدراسية: {str(e)}", "danger")
    
    return redirect(url_for('annees'))
      
@app.route('/menu')
def menu():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    if 'annee_info' not in session:
        flash('الرجاء اختيار سنة دراسية أولاً', 'warning')
        return redirect(url_for('annees'))
    
    try:
        # استيراد MENU_ITEMS كقاموس
        from menu_manager import MENU_ITEMS       
        # للتحقق: اطبع النوع والمحتوى     
        
        return render_template('menu.html',
                           
                            menu_items=MENU_ITEMS,
                            annee_info=session['annee_info'])
    except Exception as e:
        flash(f'حدث خطأ في تحميل القائمة: {str(e)}', 'danger')
        return redirect(url_for('annees'))
    
@app.route('/professeurs', methods=['GET'])
def professeurs():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    conn = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT id_em, num_em, nom_em, ader_em, tel_em, typeem, diplome FROM employeurs ORDER BY id_em DESC")
        professors = cur.fetchall()
        return render_template('professeurs.html', professors=professors)
    except Exception as e:
        flash(f"حدث خطأ: {str(e)}", "danger")
        return redirect(url_for('menu'))
    finally:
        if conn:
            conn.close()

@app.route('/add_professor', methods=['POST'])
def add_professor():
    if not session.get('logged_in'):
        return redirect(url_for('login'))    
    num_em = request.form.get('num_em')
    nom_em = request.form.get('nom_em')
    ader_em = request.form.get('ader_em')
    tel_em = request.form.get('tel_em')
    typeem = request.form.get('typeem')
    diplome = request.form.get('diplome')
    
    conn = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO employeurs (num_em, nom_em, ader_em, tel_em, typeem, diplome)
            VALUES (%s, %s, %s, %s, %s, %s)
        """, (num_em, nom_em, ader_em, tel_em, typeem, diplome))
        conn.commit()
        flash('تمت إضافة المدرس بنجاح', 'success')
    except Exception as e:
        if conn:
            conn.rollback()
        flash(f'حدث خطأ أثناء الإضافة: {str(e)}', 'danger')
    finally:
        if conn:
            conn.close()
    
    return redirect(url_for('professeurs'))

@app.route('/professeurs/update/<int:id>', methods=['POST'])
def update_professor(id):
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    print("Received form data:", request.form)  # للتحقق من البيانات المستلمة
    
    num_em = request.form.get('num_em')
    nom_em = request.form.get('nom_em')
    ader_em = request.form.get('ader_em')
    tel_em = request.form.get('tel_em')
    typeem = request.form.get('typeem')
    diplome = request.form.get('diplome')    
    conn = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            UPDATE employeurs 
            SET num_em=%s, nom_em=%s, ader_em=%s, tel_em=%s, typeem=%s, diplome=%s
            WHERE id_em=%s
            RETURNING id_em
        """, (num_em, nom_em, ader_em, tel_em, typeem, diplome, id))        
        updated_row = cur.fetchone()
        if updated_row:
            conn.commit()
            flash('تم تحديث بيانات المدرس بنجاح', 'success')
        else:
            flash('لم يتم العثور على المدرس المطلوب', 'warning')
    except Exception as e:
        if conn:
            conn.rollback()
        flash(f'حدث خطأ أثناء التحديث: {str(e)}', 'danger')
        app.logger.error(f'Error updating professor: {str(e)}')
    finally:
        if conn:
            conn.close()    
    return redirect(url_for('professeurs'))

@app.route('/professeurs/delete/<int:id>')
def delete_professor(id):
    if not session.get('logged_in'):
        return redirect(url_for('login'))    
    conn = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("DELETE FROM employeurs WHERE id_em = %s", (id,))
        conn.commit()
        flash('تم حذف المدرس بنجاح', 'success')
    except Exception as e:
        if conn:
            conn.rollback()
        flash(f'حدث خطأ أثناء الحذف: {str(e)}', 'danger')
    finally:
        if conn:
            conn.close()    
    return redirect(url_for('professeurs'))
# بقية الروتات الأخرى تبقى كما هي...
#====================================ELُ
@app.route('/eleves')
def eleves():
    if not session.get('logged_in'):
        return redirect(url_for('login'))    
    if 'annee_info' not in session:
        flash('الرجاء اختيار سنة دراسية أولاً', 'warning')
        return redirect(url_for('annees'))    
    id_ann = session['annee_info']['id_ann']    
    conn = None
    cur = None    
    try:
        conn = get_db_connection()
        cur = conn.cursor() 
        # جلب الفصول
        cur.execute('SELECT id_cl, nom_cl FROM classes')
        classes = cur.fetchall()       
        # استعلام معدل: جلب بيانات الطلاب مع اسم الصف
        cur.execute('''
            SELECT id_el, num_el, nom_el, tel, gender, ader_parent, eta_etuditn, type_el, id_cl, nom_cl  
            FROM eleves NATURAL JOIN claselves NATURAL JOIN classes 
            WHERE id_ann = %s
        ''', (id_ann,))
        eleves_list = cur.fetchall()

        cur.execute('''
            select distinct num_el, nom_el, nom_cl, type_el, montant_a_paiye, inscrption, montant_arier, 
            (montant_a_paiye+inscrption+montant_arier) as montant_total_apaiye, credit_cfaa as t_payer, 
            (montant_a_paiye+inscrption+montant_arier)-credit_cfaa as solde from eleves natural join contenu_comtable 
            natural join ann_scolaire natural join claselves natural join classes where id_ann = %s          
            
        ''', (id_ann,))
        eleves_listp = cur.fetchall()


        return render_template('eleves.html',
                            classes=classes,
                            eleves_list=eleves_list,
                            eleves_listp=eleves_listp)       
                
    except Exception as e:
        flash(f'حدث خطأ في قاعدة البيانات: {str(e)}', 'danger')
        app.logger.error(f'خطأ في قاعدة البيانات: {str(e)}')
        return redirect(url_for('menu'))
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()

@app.route('/ajouter_eleve', methods=['POST'])
def ajouter_eleve():
    if not session.get('logged_in'):
        return redirect(url_for('login'))    
    if 'annee_info' not in session:
        flash('الرجاء اختيار سنة دراسية أولاً', 'warning')
        return redirect(url_for('annees'))    
    id_ann = session['annee_info']['id_ann']
    
    if request.method == 'POST':
        conn = None
        try:
            # التحقق من الحقول المطلوبة
            required_fields = ['num_el', 'nom_el', 'tel', 'gender', 'ader_parent', 'eta_etuditn', 'type_el', 'id_cl']
            for field in required_fields:
                if field not in request.form or not request.form[field].strip():
                    flash('جميع الحقول المطلوبة يجب ملؤها', 'danger')
                    return redirect(url_for('eleves'))            
            # جمع البيانات من النموذج
            num_el = request.form['num_el'].strip()
            nom_el = request.form['nom_el'].strip()
            tel = request.form['tel'].strip()
            gender = request.form['gender']
            ader_parent = request.form['ader_parent'].strip()
            fee_type = request.form['eta_etuditn']
            student_status = request.form['type_el']
            id_cl = request.form['id_cl']            
            # تحويل الحقول الاختيارية إلى أرقام
            montant_arier = float(request.form.get('montant_arier', '0'))
            credit_cfaa = float(request.form.get('credit_cfaa', '0'))
            credit_cfa = float(request.form.get('credit_cfa', '0'))
            id_ope = request.form.get('id_ope', '0')            
            # إدخال البيانات في قاعدة البيانات
            conn = get_db_connection()
            cur = conn.cursor()            
            # 1. الحصول على رسوم الفصل والرسوم التسجيل
            cur.execute('''
                SELECT montant_a_paiye, inscrption FROM clas_frais_scol 
                NATURAL JOIN frais_scol 
                WHERE id_cl = %s AND type_frais = %s
            ''', (id_cl, fee_type))
            frais_data = cur.fetchone()
            
            if not frais_data:
                frais_classe = 0
                frais_inscription = 0
                flash('تحذير: لم يتم العثور على رسوم للفصل المحدد، تم استخدام قيم افتراضية (0)', 'warning')
            else:
                frais_classe = float(frais_data[0])
                frais_inscription = float(frais_data[1])
            
            # 2. حساب الرسوم حسب حالة الطالب
            if student_status == 'nouveau':  # طالب جديد
                montant_total = frais_classe
                frais_inscription_val = frais_inscription
            else:  # طالب قديم
                montant_total = frais_classe
                frais_inscription_val = 0            
            # 3. إدخال بيانات الطالب
            cur.execute("""
                INSERT INTO eleves (num_el, nom_el, tel, gender, ader_parent)
                VALUES (%s, %s, %s, %s, %s)
                RETURNING id_el
            """, (num_el, nom_el, tel, gender, ader_parent))
            id_el = cur.fetchone()[0]            
            # 4. إدخال بيانات فصل الطالب
            cur.execute("""
                INSERT INTO claselves 
                (id_cl, id_el, id_ann, montant_a_paiye, montant_arier, inscrption, type_el, eta_etuditn, credit_cfaa)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                id_cl, 
                id_el, 
                id_ann,
                montant_total,
                montant_arier, 
                frais_inscription_val,
                student_status,
                fee_type,
                credit_cfaa
            ))
            
            # 5. إدخال بيانات المحاسبة (تم تصحيح القوس الناقص هنا)
            cur.execute("""
                INSERT INTO contenu_comtable (id_el, id_ope, id_ann, credit_cfa)
                VALUES (%s, %s, %s, %s)
            """, (id_el, id_ope if id_ope != '0' else None, id_ann, credit_cfa))
            
            conn.commit()
            flash('تمت إضافة الطالب بنجاح', 'success')
            return redirect(url_for('eleves'))
            
        except ValueError:
            if conn:
                conn.rollback()
            flash('خطأ في القيم الرقمية المدخلة', 'danger')
            return redirect(url_for('eleves'))
        except Exception as e:
            if conn:
                conn.rollback()
            flash(f'حدث خطأ أثناء إضافة الطالب: {str(e)}', 'danger')
            return redirect(url_for('eleves'))
        finally:
            if conn:
                conn.close()   
# ... (يتبع باقي الدوال المماثلة)
@app.route('/modifier_eleve', methods=['POST'])
def modifier_eleve():
    if not session.get('logged_in'):
        return redirect(url_for('login'))    
    if 'annee_info' not in session:
        flash('الرجاء اختيار سنة دراسية أولاً', 'warning')
        return redirect(url_for('annees'))    
    id_ann = session['annee_info']['id_ann']
    conn = None
    cur = None
    try:
        # التحقق من البيانات المطلوبة
        required_fields = ['id_el', 'num_el', 'nom_el', 'gender', 'type_el', 'id_cl']
        for field in required_fields:
            if field not in request.form or not request.form[field].strip():
                flash(f'حقل {field} مطلوب', 'danger')
                return redirect(url_for('eleves'))
        # جمع البيانات من النموذج
        id_el = request.form['id_el']
        num_el = request.form['num_el']
        nom_el = request.form['nom_el']
        tel = request.form.get('tel', '')
        gender = request.form['gender']
        ader_parent = request.form.get('ader_parent', '')
        student_status = request.form['type_el']
        id_cl = request.form['id_cl']         
        # الحقول الاختياريةif student_status == 'nouveau' else None 
        fee_type = request.form.get('eta_etuditn') 
        montant_arier = request.form.get('montant_arier', '0')
        credit_cfaa = request.form.get('credit_cfaa', '0')     
        conn = get_db_connection()
        cur = conn.cursor()               
        cur.execute('''
            SELECT montant_a_paiye, inscrption 
            FROM clas_frais_scol NATURAL JOIN frais_scol 
            WHERE id_cl = %s AND type_frais = %s
        ''', (id_cl, fee_type))       
        frais_data = cur.fetchone()
        if not frais_data:
            frais_classe = 0
            frais_inscription = 0
            flash('تحذير: لم يتم العثور على رسوم للفصل المحدد، تم استخدام قيم افتراضية (0)', 'warning')
        else:
            frais_classe = frais_data[0]
            frais_inscription = frais_data[1]
                #flash('تم استخدام الرسوم المتاحة للفصل لأن النوع المحدد غير موجود', 'warning')              
            # 2. حساب الرسوم حسب حالة الطالب
        if student_status == 'nouveau':  # طالب جديد
            montant_total = frais_classe 
            #montant_total = frais_classe + frais_inscription
            frais_inscription_val = frais_inscription
        else:  # طالب قديم
            montant_total = frais_classe
            frais_inscription_val = 0
        # تحديث بيانات الطالب الأساسية
        cur.execute("""
            UPDATE eleves 
            SET num_el = %s, 
                nom_el = %s, 
                tel = %s, 
                gender = %s, 
                ader_parent = %s 
            WHERE id_el = %s
        """, (num_el, nom_el, tel, gender, ader_parent, id_el))
        # تحديث بيانات الصف
        cur.execute("""
            UPDATE claselves 
            SET id_cl = %s,
                montant_a_paiye = %s,
                montant_arier = %s,
                inscrption = %s,
                type_el = %s,
                eta_etuditn = %s,
                credit_cfaa = %s
            WHERE id_el = %s AND id_ann = %s
        """, (
            id_cl,
            montant_total,
            montant_arier,
            frais_inscription_val,
            student_status,
            fee_type,
            credit_cfaa,
            id_el,
            id_ann
        ))
        if cur.rowcount == 0:
            flash('لم يتم العثور على سجل الطالب لتحديثه', 'warning')
        else:
            conn.commit()
            flash('تم تحديث بيانات الطالب بنجاح', 'success')
    except psycopg2.Error as e:
        if conn:
            conn.rollback()
        app.logger.error(f'خطأ في قاعدة البيانات: {str(e)}')
        flash('حدث خطأ في قاعدة البيانات أثناء التحديث', 'danger')
    except Exception as e:
        if conn:
            conn.rollback()
        app.logger.error(f'خطأ غير متوقع: {str(e)}')
        flash('حدث خطأ غير متوقع أثناء التحديث', 'danger')
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()
    return redirect(url_for('eleves'))
   
@app.route('/supprimer_eleve/<id>', methods=['GET'])
def supprimer_eleve(id):
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    print("Received form data:", request.form)  # للتحقق من البيانات المستلمة
    
    conn = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("DELETE FROM eleves WHERE id_el = %s", (id,))
        conn.commit()
        flash('تم حذف الطالب بنجاح', 'success')
    except Exception as e:
        if conn:
            conn.rollback()
        flash(f'حدث خطأ أثناء الحذف: {str(e)}', 'danger')
    finally:
        if conn:
            conn.close()
    
    return redirect(url_for('eleves'))
#-------------------------------------ادخال قيود الامتحانات
@app.route('/exement')
def exement():
    if not session.get('logged_in'):
        return redirect(url_for('login'))    
    if 'annee_info' not in session:
        flash('الرجاء اختيار سنة دراسية أولاً', 'warning')
        return redirect(url_for('annees'))    
    #id_ann = session['annee_info']['id_ann']   

    id_ann = request.form.get('id_ann')
    num_ann = request.form.get('numann')
    annee_sc = request.form.get('anneesc') 
    conn = None
    cur = None    
    try:
        conn = get_db_connection()
        cur = conn.cursor() 
         # جلب الفصول
        cur.execute('''
            SELECT id_opef, date_opef, no_opef, nom_periode, id_annf 
            FROM operation 
            ''')        
        examents = cur.fetchall()
        return render_template('exement.html', examents=examents)                  
                            
    except Exception as e:
        flash(f'حدث خطأ في قاعدة البيانات: {str(e)}', 'danger')
        app.logger.error(f'خطأ في قاعدة البيانات: {str(e)}')
        return redirect(url_for('menu'))
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()

@app.route('/ajouter_operation', methods=['POST'])
def ajouter_operation():
    if not session.get('logged_in'):
        return redirect(url_for('login'))    
    if 'annee_info' not in session:
        flash('الرجاء اختيار سنة دراسية أولاً', 'warning')
        return redirect(url_for('annees'))

    if request.method == 'POST':
        conn = None
        try:
            # جلب البيانات من النموذج
            date_opef = request.form['date_opef']
            no_opef = request.form['no_opef']
            nom_periode = request.form['nom_periode']            
            id_opef = request.form.get('id_opef')
            id_ann = request.form.get('id_ann') or session.get('annee_info', {}).get('id_ann')
            # الحصول على id_annf من الجلسة مع قيمة افتراضية إذا لم تكن موجودة              
            if not id_ann:
                flash('لم يتم تحديد السنة الدراسية', 'danger')
                return redirect(url_for('exement'))            
            # تحويل تنسيق التاريخ إذا لزم الأمر
            conn = get_db_connection()
            cur = conn.cursor()            
            if id_ann:  # عملية التعديل            
                cur.execute("""
                    INSERT INTO operation (date_opef, no_opef, nom_periode, id_annf)
                    VALUES (%s, %s, %s, %s)
                    RETURNING id_opef
                """, (date_opef, no_opef, nom_periode, id_ann))
                flash('تمت إضافة البيانات بنجاح', 'success')
            
            conn.commit()
            return redirect(url_for('exement'))
            
        except Exception as e:
            if conn:
                conn.rollback()
            flash(f'حدث خطأ أثناء العملية: {str(e)}', 'danger')
            return redirect(url_for('exement'))
        finally:
            if conn:
                conn.close()

@app.route('/modifier_exement', methods=['POST'])
def modifier_exement():
    if not session.get('logged_in'):
        return redirect(url_for('login'))    
    if 'annee_info' not in session:
        flash('الرجاء اختيار سنة دراسية أولاً', 'warning')
        return redirect(url_for('annees'))

    if request.method == 'POST':
        conn = None
        try:
            id_opef = request.form['id_opef']
            date_opef = request.form['date_opef']
            no_opef = request.form['no_opef']
            nom_periode = request.form['nom_periode']
            id_annf = request.form['id_annf']
            
            conn = get_db_connection()
            cur = conn.cursor()
            
            cur.execute("""
                UPDATE operation 
                SET date_opef=%s, no_opef=%s, nom_periode=%s, id_annf=%s
                WHERE id_opef=%s
            """, (date_opef, no_opef, nom_periode, id_annf, id_opef))
            
            conn.commit()
            flash('تم تحديث البيانات بنجاح', 'success')
            return redirect(url_for('exement'))
            
        except Exception as e:
            if conn:
                conn.rollback()
            flash(f'حدث خطأ أثناء التحديث: {str(e)}', 'danger')
            return redirect(url_for('exement'))
        finally:
            if conn:
                conn.close()

@app.route('/delete_exement/<id>', methods=['GET'])
def delete_exement(id):
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    print("Received form data:", request.form)  # للتحقق من البيانات المستلمة
    
    conn = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("DELETE FROM operation WHERE id_opef = %s", (id,))
        conn.commit()
        flash('تم حذف البيانات بنجاح', 'success')
    except Exception as e:
        if conn:
            conn.rollback()
        flash(f'حدث خطأ أثناء الحذف: {str(e)}', 'danger')
    finally:
        if conn:
            conn.close()
    
    return redirect(url_for('exement'))

#=======================================================
# --- المرحلة 1: اختيار رقم قيد الامتحان ---
@app.route('/select_operation', methods=['GET', 'POST'])
def select_operation():
    #print("----- بدء select_operation -----")  # طباعة للتتبع
    #print("method:", request.method)  # نوع الطلب
    
    if not session.get('logged_in'):
        print("لم يتم تسجيل الدخول")
        return redirect(url_for('login'))
    
    if 'annee_info' not in session:
        #print("لم يتم اختيار سنة دراسية")
        return redirect(url_for('annees'))

    if request.method == 'GET':
        #print("معالجة طلب GET")
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute('SELECT id_opef, nom_periode, id_annf FROM operation ORDER BY id_opef')
            operations = cur.fetchall()
            print("عدد العمليات المسترجعة:", len(operations))  # طباعة عدد العمليات
            
            return render_template('select_operation.html', 
                                operations=operations,
                                current_annee=session.get('annee_info'))
        except Exception as e:
            #print("حدث خطأ في GET:", str(e))  # طباعة الخطأ
            flash(f'حدث خطأ في جلب العمليات: {str(e)}', 'danger')
            return redirect(url_for('dashboard'))
        finally:
            if 'cur' in locals(): cur.close()
            if 'conn' in locals(): conn.close()

    elif request.method == 'POST':
        print("معالجة طلب POST")
        selected_opef_id = request.form.get('id_opef')
        print("ID العملية المختارة:", selected_opef_id)  # طباعة ID العملية
        
        if not selected_opef_id:
            print("لم يتم اختيار عملية")
            flash('الرجاء اختيار عملية', 'danger')
            return redirect(url_for('select_operation'))
        
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            
            #print("تنفيذ استعلام العملية...")
            cur.execute('SELECT id_opef, nom_periode FROM operation WHERE id_opef = %s', (selected_opef_id,))
            op_data = cur.fetchone()
            #print("نتيجة الاستعلام:", op_data)  # طباعة نتيجة الاستعلام
            
            if op_data:
                session['selected_opef'] = op_data[0]
                session['selected_opef_name'] = op_data[1]
                #print("تم تخزين في الجلسة:", session['selected_opef'], session['selected_opef_name'])
                
                #print("التوجيه إلى select_class...")
                return redirect(url_for('select_class'))
            else:
                #print("العملية غير موجودة في DB")
                flash('العملية المحددة غير موجودة', 'danger')
        except Exception as e:
            #print("حدث خطأ في POST:", str(e))
            flash(f'حدث خطأ: {str(e)}', 'danger')
        finally:
            if 'cur' in locals(): cur.close()
            if 'conn' in locals(): conn.close()
        
        return redirect(url_for('select_operation'))

@app.route('/select_class', methods=['GET', 'POST'])
def select_class():
    if 'selected_opef' not in session:
        flash('الرجاء اختيار عملية أولاً', 'warning')
        return redirect(url_for('select_operation'))
    
    conn = None
    cur = None
    
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        
        cur.execute('''
            SELECT c.id_cl, c.nom_cl, s.nom_se, s.id_se
            FROM classes c
            JOIN series s ON c.id_se = s.id_se
            ORDER BY c.nom_cl
        ''')
        classes = cur.fetchall()
        
        if request.method == 'POST':
            if 'id_cl' not in request.form:
                flash('الرجاء اختيار فصل', 'danger')
                return redirect(url_for('select_class'))
            
            try:
                selected_cl_id = int(request.form['id_cl'])
                selected_class = next((c for c in classes if c[0] == selected_cl_id), None)
                
                if selected_class:
                    session['selected_cl'] = selected_cl_id
                    session['selected_cl_name'] = selected_class[1]
                    session['selected_cl_series'] = selected_class[3]  # حفظ id_se في الجلسة
                    
                    # تحسين منطق التوجيه
                    if selected_class[3] in (115, 120):  # أو أي قيم أخرى تحتاجها
                        return redirect(url_for('add_notes'))
                    else:
                        return redirect(url_for('add_notes2'))
                else:
                    flash('الفصل المحدد غير صالح', 'danger')
            except ValueError:
                flash('معرف الفصل غير صحيح', 'danger')
            
            return redirect(url_for('select_class'))
        
        return render_template('select_class.html',
                            classes=classes,
                            operation_name=session.get('selected_opef_name', 'عملية غير معروفة'))
    
    except Exception as e:
        flash(f'حدث خطأ: {str(e)}', 'danger')
        return redirect(url_for('select_operation'))
    finally:
        if cur: cur.close()
        if conn: conn.close()


@app.route('/add_notes', methods=['GET', 'POST'])
def add_notes():
    # تحقق إضافي من سلسلة الفصل
    if 'selected_cl_series' not in session or session['selected_cl_series'] not in (115, 120):
        flash('هذه الصفحة غير متاحة لهذا الفصل', 'warning')
        return redirect(url_for('select_class'))
    
    return handle_notes_template('add_notes.html')


@app.route('/add_notes2', methods=['GET', 'POST'])
def add_notes2():
    return handle_notes_template('add_notes2.html')

def handle_notes_template(template_name):
    # التحقق من وجود العملية والفصل المحدد
    if 'selected_opef' not in session or 'selected_cl' not in session:
        flash('الرجاء اختيار العملية والفصل أولاً', 'warning')
        return redirect(url_for('select_operation'))
    
    # جلب البيانات من الجلسة
    id_opef = session['selected_opef']
    id_cl = session['selected_cl']
    id_ann = session.get('selected_annee', 151)  # قيمة افتراضية إذا لم تكن موجودة
    op_name = session.get('selected_opef_name', 'عملية غير معروفة')
    cl_name = session.get('selected_cl_name', 'فصل غير معروف')
    
    conn = get_db_connection()
    cur = conn.cursor()
    
    try:
        # جلب المواد الخاصة بالفصل مع معاملاتها
        cur.execute('''
            SELECT id_ma, nom_ma, coef_ma FROM matieres NATURAL JOIN matiereclasse NATURAL JOIN classes 
            WHERE id_cl = %s
            ORDER BY nom_ma
        ''', (id_cl,))
        matieres = cur.fetchall()

        # جلب طلاب الفصل
        cur.execute('''
            SELECT id_el, nom_el, num_el FROM eleves NATURAL JOIN claselves 
            WHERE id_cl = %s
            ORDER BY nom_el
        ''', (id_cl,))    
        eleves = cur.fetchall()

        cur.execute('''
            SELECT id_ne_cl, num_el, nom_el, nom_ma, round(nolesel,2), round(note_compo,2), 
                   round(moyen,2), round(moy_coeff,2), etat_ma, id_ope 
            FROM noteesfr 
            NATURAL JOIN eleves 
            NATURAL JOIN matieres 
            NATURAL JOIN classes 
            NATURAL JOIN claselves 
            WHERE id_cl=%s AND id_ope=%s           
            ORDER BY id_ne_cl desc
        ''', (id_cl, id_opef))    
        notescf = cur.fetchall()

        cur.execute('''
            SELECT num_el, nom_el, round(som_notes_period, 2), etat_notes_period 
            FROM eleve_note 
            NATURAL JOIN eleves 
            NATURAL JOIN classes 
            NATURAL JOIN claselves
            WHERE id_ope=%s AND id_cl=%s
        ''', (id_opef, id_cl))    
        notescf2 = cur.fetchall() 

        # جلب درجات المواد الحالية
        cur.execute('''
            SELECT id_el, id_ma, nolesel, note_compo, moy_coeff, etat_ma
            FROM noteesfr WHERE id_ope = %s AND id_el IN (SELECT id_el FROM claselves WHERE id_cl = %s)
        ''', (id_opef, id_cl))
        existing_notes = {f"{row[0]}_{row[1]}": row[2:] for row in cur.fetchall()}

        # جلب ملخص درجات الطلاب
        cur.execute('''
            SELECT id_el, som_notes_period, etat_notes_period FROM eleve_note
            WHERE id_ope = %s AND id_el IN (SELECT id_el FROM claselves WHERE id_cl = %s)
        ''', (id_opef, id_cl))
        student_summaries = {row[0]: row[1:] for row in cur.fetchall()}

        # معالجة إرسال النموذج
        if request.method == 'POST':
            try:
                id_el = int(request.form['id_el'])
                id_ma = int(request.form['id_ma'])
                nolesel = float(request.form.get('nolesel', 0))
                note_compo = float(request.form.get('note_compo', 0))
                
                # التحقق من صحة الدرجة
                if not (0 <= nolesel <= 20) or not (0 <= note_compo <= 20):
                    flash('الدرجة يجب أن تكون بين 0 و 20', 'danger')
                    return redirect(request.url)

                # الحصول على معامل المادة
                coef_ma = next((m[2] for m in matieres if m[0] == id_ma), 1.0)
                moy_coeff = float(note_compo) 
                moyen = float(note_compo)  # يمكن تعديله حسب احتياجاتك* float(coef_ma) 

                # التحقق من وجود الدرجة
                cur.execute('''
                    SELECT 1 FROM noteesfr 
                    WHERE id_el = %s AND id_ma = %s AND id_ope = %s
                ''', (id_el, id_ma, id_opef))
                exists = cur.fetchone()

                if exists:
                    # تحديث الدرجة
                    cur.execute('''
                        UPDATE noteesfr 
                        SET note_compo = %s, 
                            nolesel = %s,
                            moyen = %s,
                            moy_coeff = %s
                            
                        WHERE id_el = %s AND id_ma = %s AND id_ope = %s
                    ''', (note_compo, nolesel, moyen, moy_coeff, id_el, id_ma, id_opef))
                    flash('تم تحديث الدرجة بنجاح', 'success')
                else:
                    # إضافة درجة جديدة
                    cur.execute('''
                        INSERT INTO noteesfr 
                        (id_el, nolesel, id_ma, id_ope, etat_ma, moyen, moy_coeff, note_compo)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                    ''', (id_el, nolesel, id_ma, id_opef, 'non', moyen, moy_coeff, note_compo))
                    flash('تمت إضافة الدرجة بنجاح', 'success')

                # حساب عدد المواد والطلاب
                cur.execute('''
                    SELECT COUNT(DISTINCT id_ma) 
                    FROM matiereclasse 
                    WHERE id_cl = %s
                ''', (id_cl,))
                num_subjects = cur.fetchone()[0] or 0

                cur.execute('''
                    SELECT COUNT(DISTINCT id_el) 
                    FROM claselves 
                    WHERE id_cl = %s
                ''', (id_cl,))
                num_students = cur.fetchone()[0] or 0

                expected_records = num_students * num_subjects

                cur.execute('''
                    SELECT COUNT(*) 
                    FROM noteesfr 
                    WHERE id_ope = %s AND id_el IN (
                        SELECT id_el FROM claselves WHERE id_cl = %s
                    )
                ''', (id_opef, id_cl))
                actual_records = cur.fetchone()[0] or 0

                # إذا اكتملت جميع الدرجات
                if actual_records == expected_records and expected_records > 0:
                    # حساب المعدل لكل طالب
                    cur.execute('''
                        SELECT id_el, ROUND(SUM(moy_coeff) / SUM(coef_ma), 2) FROM noteesfr NATURAL JOIN matiereclasse NATURAL JOIN matieres 
                        WHERE id_cl = %s AND id_ope = %s GROUP BY id_el
                    ''', (id_cl, id_opef))
                    
                    for student in cur.fetchall():
                        id_el, moyenne = student
                        # تحديد حالة الطالب
                        if moyenne >= 17.99:
                            etat = 'ممتاز (Excellent)'
                        elif moyenne >= 15.99:
                            etat = 'جيد جدا (Très Bien)'
                        elif moyenne >= 13.99:
                            etat = 'جيد (Bien)'
                        elif moyenne >= 12.00:
                            etat = 'مستحسن (Assez Bien)'
                        elif moyenne >= 10.00:
                            etat = 'مقبول (Passable)'
                        elif moyenne >= 8.00:
                            etat = 'غير كاف (Insuffisant)'
                        elif moyenne >= 6.00:
                            etat = 'ضعيف (Faible)'
                        elif moyenne >= 4.00:
                            etat = 'ضعيف جدا (Très Faible)'
                        else:
                            etat = 'راسب (Ajourné)'
                      
                        # أولاً: محاولة التحديث
                        update_query = """
                            UPDATE eleve_note 
                            SET som_notes_period = %s,
                                etat_notes_period = %s,
                                id_ann = %s
                            WHERE id_el = %s AND id_ope = %s
                        """
                        cur.execute(update_query, (moyenne, etat, id_ann, id_el, id_opef))

                        # إذا لم يتم تحديث أي صف (يعني السجل غير موجود)
                        if cur.rowcount == 0:
                            insert_query = """
                                INSERT INTO eleve_note 
                                (id_el, id_ope, som_notes_period, etat_notes_period, id_ann)
                                VALUES (%s, %s, %s, %s, %s)
                            """
                            cur.execute(insert_query, (id_el, id_opef, moyenne, etat, id_ann))

                conn.commit()
                updaetamacf()
                inserreng()                
                return redirect(request.url)             

            except Exception as e:
                conn.rollback()
                flash(f'خطأ في حفظ البيانات: {str(e)}', 'danger')
                import traceback
                traceback.print_exc()
                return redirect(request.url)
        # عرض القالب
        return render_template(template_name, 
                            matieres=matieres,
                            eleves=eleves,
                            existing_notes=existing_notes,
                            student_summaries=student_summaries,
                            operation_name=op_name,
                            class_name=cl_name,
                            operation_id=id_opef,
                            class_id=id_cl,
                            notescf=notescf, 
                            notescf2=notescf2)

    except Exception as e:
        conn.rollback()
        flash(f'حدث خطأ: {str(e)}', 'danger')
        return redirect(url_for('select_class'))
    finally:
        if cur: cur.close()
        if conn: conn.close()

@app.route('/updaetamacf')
def updaetamacf():
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cursor = conn.cursor()        
        # 1. جلب آخر سجل
        cursor.execute("SELECT id_ne_cl, id_ope, id_el FROM noteesfr ORDER BY id_ne_cl DESC LIMIT 1")
        last_record = cursor.fetchone()        
        if not last_record:
            flash('لا توجد بيانات مسجلة', 'info')
            return redirect(url_for('some_route'))        
        id_ne_cl, id_ope, id_el = last_record
        # 2. جلب بيانات الطالب
        query = """
            SELECT DISTINCT id_ne_cl, id_el, id_ope, note_compo, nolesel,
                CASE 
                    WHEN id_ma != 473 THEN  
                        ROUND(((note_compo * 2) + nolesel) / 3, 2)
                    ELSE
                        nolesel * 1  
                END AS moyene, 
                coef_ma 
            FROM noteesfr 
            NATURAL JOIN matieres 
            NATURAL JOIN matiereclasse 
            WHERE id_ne_cl = %s  
        """
        cursor.execute(query, (id_ne_cl,))
        records = cursor.fetchall()
        if not records:
            flash('لا توجد بيانات مسجلة', 'info')
            return redirect(url_for('some_route'))

        results = []
        for record in records:
            # 3. حساب القيم
            moyen_val = record[5]
            coef_val = record[6]
            moycoeff_val = round(moyen_val * coef_val, 2)
            # 4. تحديد الحالة
            if moycoeff_val >= 17.99:
                etat_ma = 'ممتاز (Excellent)'
            elif moycoeff_val >= 15.99:
                etat_ma = 'جيد جدا (Très Bien)'
            elif moycoeff_val >= 13.99:
                etat_ma = 'جيد(Bien)'
            elif moycoeff_val >= 12.99:
                etat_ma = 'مستحسن(Assez Bien)'
            elif moycoeff_val >= 10.99:
                etat_ma = 'مقبول(Passable)'
            elif moycoeff_val >= 8.99:
                etat_ma = 'غير كاف(Insuffisant)'
            elif moycoeff_val >= 6.99:
                etat_ma = 'ضعيف(Faible)'
            elif moycoeff_val >= 4.99:
                etat_ma = 'ضعيف جدا(Très Faible)'
            else:
                etat_ma = 'لا شيء(Nul)'

            # 5. تحديث قاعدة البيانات
            update_query = """
                UPDATE noteesfr 
                SET etat_ma = %s,
                    moyen = %s,
                    moy_coeff = %s
                WHERE id_ne_cl = %s
                  AND id_el = %s
                  AND id_ope = %s
            """
            cursor.execute(update_query, (
                etat_ma,
                moyen_val,
                moycoeff_val,
                record[0],  # id_ne_cl
                record[1],  # id_el
                record[2]   # id_ope
            ))
            
            results.append({
                'id_ne_cl': record[0],
                'id_el': record[1],
                'id_ope': record[2],
                'note_compo': record[3],
                'nolesel': record[4],
                'moyen': moyen_val,
                'coef_ma': coef_val,
                'moy_coeff': moycoeff_val,
                'etat_ma': etat_ma
            })

        conn.commit()
        #flash('تم تحديث البيانات بنجاح', 'success')
        return jsonify({'status': 'success', 'data': results})

    except Exception as e:
        conn.rollback()
        flash(f'حدث خطأ: {str(e)}', 'danger')
        return jsonify({'status': 'error', 'message': str(e)}), 500

    finally:
        if 'cursor' in locals():
            cursor.close()
        if 'conn' in locals():
            conn.close()

@app.route('/inserreng')
def inserreng():
    try:        
        id_opef = session['selected_opef']
        id_cl = session['selected_cl']
        id_ann = session.get('selected_annee', 151)
        
        if not id_cl or not id_opef:
            flash('يجب إدخال معرف الفصل والعملية', 'danger')
            return False

        conn = psycopg2.connect(**DB_CONFIG)
        cursor = conn.cursor()

        # التحقق من اكتمال الدرجات
        cursor.execute("SELECT COUNT(DISTINCT id_ma) FROM matiereclasse WHERE id_cl = %s", (id_cl,))
        num_subjects = cursor.fetchone()[0] or 0

        cursor.execute("SELECT COUNT(DISTINCT id_el) FROM claselves WHERE id_cl = %s", (id_cl,))
        num_students = cursor.fetchone()[0] or 0

        expected_records = num_students * num_subjects

        cursor.execute("""
            SELECT COUNT(*) FROM noteesfr 
            WHERE id_ope = %s AND id_el IN (
                SELECT id_el FROM claselves WHERE id_cl = %s
            )
        """, (id_opef, id_cl))
        actual_records = cursor.fetchone()[0] or 0

        if actual_records != expected_records or expected_records == 0:
            flash(f'لم يتم إدخال جميع الدرجات بعد. المطلوب: {expected_records}، الموجود: {actual_records}', 'warning')
            return False

        # حساب الترتيب والمعدل السنوي
        cursor.execute("""
            SELECT id_el, id_ope, nom_el, som_notes_period 
            FROM eleve_note 
            NATURAL JOIN eleves 
            NATURAL JOIN claselves 
            WHERE id_cl = %s AND id_ope = %s
            ORDER BY som_notes_period DESC
        """, (id_cl, id_opef))

        students = cursor.fetchall()
        current_rank = 1
        previous_score = None

        for idx, (id_el, id_ope, nom_el, score) in enumerate(students, start=1):
            if previous_score is not None and score != previous_score:
                current_rank = idx
                
            # تحديث ترتيب الطالب للفترة الحالية
            cursor.execute("""
                UPDATE eleve_note
                SET ranking = %s
                WHERE id_el = %s AND id_ope = %s
            """, (current_rank, id_el, id_opef))
            
            previous_score = score

            # حساب المعدل السنوي (متوسط جميع الفترات)
            cursor.execute("""
                select SUM(som_notes_period) / (COUNT(DISTINCT id_ma) * COUNT(DISTINCT id_ope)) AS moyenne_alternat 
                FROM eleve_note natural join eleves natural join classes natural join claselves natural join matiereclasse                
                WHERE id_el = %s AND id_ann = %s GROUP BY id_el
    
            """, (id_el, id_ann))
            annual_avg = cursor.fetchone()[0]

            # تحديث المعدل السنوي في جدول claselves للفترة الحالية
            cursor.execute("""
                UPDATE claselves
                SET moyanuel = %s
                WHERE id_el = %s AND id_ann = %s AND id_cl = %s 
            """, (annual_avg, id_el, id_ann, id_cl))

        conn.commit()
        flash('تم تحديث ترتيب الطلاب والمعدل السنوي بنجاح', 'success')
        return True

    except psycopg2.Error as e:
        conn.rollback()
        flash(f'خطأ في قاعدة البيانات: {str(e)}', 'danger')
        return False
        
    except Exception as e:
        flash(f'حدث خطأ غير متوقع: {str(e)}', 'danger')
        return False
        
    finally:
        if 'cursor' in locals():
            cursor.close()
        if 'conn' in locals():
            conn.close()
     
@app.route('/modifier_notes', methods=['POST'])
def modifier_notes():
    conn = None
    try:
        # التحقق من الحقول المطلوبة
        required_fields = {
            'id_ne_cl': int,
            'id_el': int,
            'id_ma': int,
            'id_ope': int,
            'nolesel': float,
            'note_compo': float,
            'moyen': float,
            'moy_coeff': float
        }
        
        id_ann = session.get('selected_annee', 151)  # قيمة افتراضية
        data = {}        
        for field, field_type in required_fields.items():
            value = request.form.get(field)
            if not value:
                flash(f'حقل {field} مطلوب', 'danger')
                return redirect(url_for('add_notes'))

            try:
                if field_type == float:
                    data[field] = round(float(value), 2)
                else:
                    data[field] = int(value)
            except ValueError:
                flash(f'قيمة غير صالحة في حقل {field}، يجب أن تكون رقمية', 'danger')
                return redirect(url_for('add_notes'))

        if any(v < 0 for v in [data['nolesel'], data['note_compo'], data['moyen'], data['moy_coeff']]):
            flash('القيم لا يمكن أن تكون سالبة', 'danger')
            return redirect(url_for('add_notes'))

        conn = get_db_connection()
        cur = conn.cursor()        
        # التحديث الأساسي مع التحقق من القيم الرقمية
        try:
            cur.execute("""
                UPDATE noteesfr 
                SET id_el = %s, 
                    nolesel = %s, 
                    id_ma = %s, 
                    id_ope = %s, 
                    note_compo = %s,
                    moyen = %s,
                    moy_coeff = %s
                WHERE id_ne_cl = %s
                RETURNING id_ne_cl
            """, (
                data['id_el'], 
                data['nolesel'], 
                data['id_ma'], 
                data['id_ope'], 
                data['note_compo'],
                data['moyen'], 
                data['moy_coeff'], 
                data['id_ne_cl']
            ))

            if not cur.fetchone():
                flash('لم يتم العثور على السجل للتحديث', 'warning')
                conn.close()
                return redirect(url_for('add_notes'))
        except psycopg2.Error as e:
            flash(f'خطأ في تحديث السجل الأساسي: {str(e)}', 'danger')
            conn.rollback()
            return redirect(url_for('add_notes'))
        # حساب المعدل العام لكل مادة
        try:
            query = """
                SELECT DISTINCT 
                    id_ne_cl, 
                    id_el, 
                    id_ope, 
                    CAST(note_compo AS FLOAT), 
                    CAST(nolesel AS FLOAT),
                    CASE 
                        WHEN id_ma != 473 THEN  
                            ROUND(((note_compo * 2) + nolesel) / 3, 2)
                        ELSE
                            nolesel * 1  
                    END AS moyene, 
                    CAST(coef_ma AS FLOAT)
                FROM noteesfr 
                NATURAL JOIN matieres 
                NATURAL JOIN matiereclasse
                WHERE id_ne_cl = %s  
            """
            cur.execute(query, (data['id_ne_cl'],))
            matieres_rows = cur.fetchall()

            for row in matieres_rows:
                try:
                    id_ne_cl, id_el, id_ope, note_compo, nolesel, moyene, coef_ma = row                    
                    # التحويل الآمن للقيم العددية
                    try:
                        moyene = float(moyene) if moyene is not None else 0.0
                        coef_ma = float(coef_ma) if coef_ma is not None else 0.0
                        moy_coeff = round(moyene * coef_ma, 2)
                    except (ValueError, TypeError) as e:
                        flash(f'قيمة غير صالحة في سجل المادة {id_ne_cl}: {str(e)}', 'warning')
                        continue
                    # تحديد التقييم النصي
                    if moy_coeff >= 17.99:
                        etat_ma = 'ممتاز (Excellent)'
                    elif moy_coeff >= 15.99:
                        etat_ma = 'جيد جدا (Très Bien)'
                    elif moy_coeff >= 13.99:
                        etat_ma = 'جيد(Bien)'
                    elif moy_coeff >= 12.99:
                        etat_ma = 'مستحسن(Assez Bien)'
                    elif moy_coeff >= 10.99:
                        etat_ma = 'مقبول(Passable)'
                    elif moy_coeff >= 8.99:
                        etat_ma = 'غير كاف(Insuffisant)'
                    elif moy_coeff >= 6.99:
                        etat_ma = 'ضعيف(Faible)'
                    elif moy_coeff >= 4.99:
                        etat_ma = 'ضعيف جدا(Très Faible)'
                    else:
                        etat_ma = 'لا شيء(Nul)'

                    # تحديث المادة
                    cur.execute("""
                        UPDATE noteesfr 
                        SET etat_ma = %s,
                            moyen = %s,
                            moy_coeff = %s
                        WHERE id_ne_cl = %s
                          AND id_el = %s
                          AND id_ope = %s
                    """, (
                        etat_ma,
                        moyene,
                        moy_coeff,
                        id_ne_cl,
                        id_el,
                        id_ope
                    ))

                except Exception as e:
                    flash(f'خطأ في معالجة مادة الطالب {id_el}: {str(e)}', 'warning')
                    conn.rollback()
                    continue
            # حساب المعدل العام للطالب
            query = """        
                SELECT 
                    id_el, 
                    id_ope, 
                    ROUND(SUM(moy_coeff) / NULLIF(SUM(coef_ma), 0), 2) AS moyen         
                FROM noteesfr 
                NATURAL JOIN matiereclasse 
                NATURAL JOIN matieres        
                WHERE id_ne_cl = %s         
                GROUP BY id_el, id_ope        
            """        
            cur.execute(query, (data['id_ne_cl'],))        
            moy_general_rows = cur.fetchall()
            for row in moy_general_rows:
                try:
                    id_el, id_ope, moyen = row
                    moyen = float(moyen) if moyen is not None else 0.0                    
                    # تحديد الحالة بناء على المعدل العام
                    if moyen >= 17.99:
                        etat_ma = 'ممتاز (Excellent)'
                    elif moyen >= 15.99:
                        etat_ma = 'جيد جدا (Très Bien)'
                    elif moyen >= 13.99:
                        etat_ma = 'جيد(Bien)'
                    elif moyen >= 12.99:
                        etat_ma = 'مستحسن(Assez Bien)'
                    elif moyen >= 10.99:
                        etat_ma = 'مقبول(Passable)'
                    elif moyen >= 8.99:
                        etat_ma = 'غير كاف(Insuffisant)'
                    elif moyen >= 6.99:
                        etat_ma = 'ضعيف(Faible)'
                    elif moyen >= 4.99:
                        etat_ma = 'ضعيف جدا(Très Faible)'
                    else:
                        etat_ma = 'لا شيء(Nul)'

                    # تحديث المعدل العام
                    cur.execute("""
                        UPDATE eleve_note 
                        SET som_notes_period = %s,
                            etat_notes_period = %s,
                            id_ann = %s
                        WHERE id_el = %s AND id_ope = %s
                    """, (
                        moyen, 
                        etat_ma, 
                        id_ann, 
                        id_el, 
                        id_ope
                    ))
                except Exception as e:
                    flash(f'خطأ في تحديث المعدل العام للطالب {id_el}: {str(e)}', 'warning')
                    conn.rollback()
                    continue
            
            conn.commit()
            inserreng2()
            flash('تم تحديث البيانات بنجاح', 'success')

        except Exception as e:
            conn.rollback()
            flash(f'حدث خطأ غير متوقع أثناء معالجة البيانات: {str(e)}', 'danger')            
    except psycopg2.Error as e:
        if conn:
            conn.rollback()
        flash(f'خطأ في قاعدة البيانات: {str(e)}', 'danger')
    except Exception as e:
        if conn:
            conn.rollback()
        flash(f'حدث خطأ غير متوقع: {str(e)}', 'danger')
    finally:
        if conn:
            conn.close()
    return redirect(url_for('add_notes'))


@app.route('/inserreng2')
def inserreng2():
    try:        
        id_opef = session['selected_opef']
        id_cl = session['selected_cl']
        id_ann = session.get('selected_annee', 151)
        
        if not id_cl or not id_opef:
            flash('يجب إدخال معرف الفصل والعملية', 'danger')
            return False

        conn = psycopg2.connect(**DB_CONFIG)
        cursor = conn.cursor()

        # التحقق من اكتمال الدرجات
        cursor.execute("SELECT COUNT(DISTINCT id_ma) FROM matiereclasse WHERE id_cl = %s", (id_cl,))
        num_subjects = cursor.fetchone()[0] or 0

        cursor.execute("SELECT COUNT(DISTINCT id_el) FROM claselves WHERE id_cl = %s", (id_cl,))
        num_students = cursor.fetchone()[0] or 0

        expected_records = num_students * num_subjects

        cursor.execute("""
            SELECT COUNT(*) FROM noteesfr 
            WHERE id_ope = %s AND id_el IN (
                SELECT id_el FROM claselves WHERE id_cl = %s
            )
        """, (id_opef, id_cl))
        actual_records = cursor.fetchone()[0] or 0

        if actual_records != expected_records or expected_records == 0:
            flash(f'لم يتم إدخال جميع الدرجات بعد. المطلوب: {expected_records}، الموجود: {actual_records}', 'warning')
            return False

        # حساب الترتيب والمعدل السنوي
        cursor.execute("""
            SELECT id_el, id_ope, nom_el, som_notes_period 
            FROM eleve_note 
            NATURAL JOIN eleves 
            NATURAL JOIN claselves 
            WHERE id_cl = %s AND id_ope = %s
            ORDER BY som_notes_period DESC
        """, (id_cl, id_opef))

        students = cursor.fetchall()
        current_rank = 1
        previous_score = None

        for idx, (id_el, id_ope, nom_el, score) in enumerate(students, start=1):
            if previous_score is not None and score != previous_score:
                current_rank = idx
                
            # تحديث ترتيب الطالب للفترة الحالية
            cursor.execute("""
                UPDATE eleve_note
                SET ranking = %s
                WHERE id_el = %s AND id_ope = %s
            """, (current_rank, id_el, id_opef))
            
            previous_score = score

            # حساب المعدل السنوي (متوسط جميع الفترات)
            cursor.execute("""
                select SUM(som_notes_period) / (COUNT(DISTINCT id_ma) * COUNT(DISTINCT id_ope)) AS moyenne_alternat 
                FROM eleve_note natural join eleves natural join classes natural join claselves natural join matiereclasse                
                WHERE id_el = %s AND id_ann = %s GROUP BY id_el
    
            """, (id_el, id_ann))
            annual_avg = cursor.fetchone()[0]

            # تحديث المعدل السنوي في جدول claselves للفترة الحالية
            cursor.execute("""
                UPDATE claselves
                SET moyanuel = %s
                WHERE id_el = %s AND id_ann = %s AND id_cl = %s 
            """, (annual_avg, id_el, id_ann, id_cl))

        conn.commit()
        flash('تم تحديث ترتيب الطلاب والمعدل السنوي بنجاح', 'success')
        return True

    except psycopg2.Error as e:
        conn.rollback()
        flash(f'خطأ في قاعدة البيانات: {str(e)}', 'danger')
        return False
        
    except Exception as e:
        flash(f'حدث خطأ غير متوقع: {str(e)}', 'danger')
        return False
        
    finally:
        if 'cursor' in locals():
            cursor.close()
        if 'conn' in locals():
            conn.close()

@app.route('/updatemanuel')
def updatemanuel():
    try:
        # جلب البيانات من النموذج (افترضنا استخدام request.form في Flask)
        #idclreh = request.form.get('idclreh')
        #idoperchx = request.form.get('idoperchx')

        id_opef = session['selected_opef']
        id_cl = session['selected_cl']
        id_ann = session.get('selected_annee', 151)  # قيمة افتراضية إذا لم تكن موجودة
        
        if not id_cl or not id_opef:
            flash('يجب إدخال معرف الفصل والعملية', 'danger')
            return False

        conn = psycopg2.connect(**DB_CONFIG)
        cursor = conn.cursor()

        # 1. حساب عدد المواد في الفصل
        cursor.execute("""
            SELECT COUNT(DISTINCT id_ma) FROM matiereclasse WHERE id_cl = %s
        """, (id_cl,))
        num_subjects = cursor.fetchone()[0] or 0

        # 2. حساب عدد الطلاب في الفصل
        cursor.execute("""
            SELECT COUNT(DISTINCT id_el) FROM claselves WHERE id_cl = %s
        """, (id_cl,))
        num_students = cursor.fetchone()[0] or 0

        # 3. حساب العدد المتوقع للسجلات
        expected_records = num_students * num_subjects

        # 4. حساب العدد الفعلي للسجلات
        cursor.execute("""
            SELECT COUNT(*) FROM noteesfr 
            WHERE id_ope = %s AND id_el IN (
                SELECT id_el FROM claselves WHERE id_cl = %s
            )
        """, (id_opef, id_cl))
        actual_records = cursor.fetchone()[0] or 0

        if actual_records != expected_records or expected_records == 0:
            #flash(f'لم يتم إدخال جميع الدرجات بعد. المطلوب: {expected_records}، الموجود: {actual_records}', 'warning')
            return False

        # 5. حساب الترتيب
        cursor.execute("""
            SELECT id_el, id_ope, nom_el, som_notes_period 
            FROM eleve_note 
            NATURAL JOIN eleves 
            NATURAL JOIN claselves 
            WHERE id_cl = %s AND id_ope = %s
            ORDER BY som_notes_period DESC
        """, (id_cl, id_opef))

        students = cursor.fetchall()
        current_rank = 1
        previous_score = None

        for idx, (id_el, id_ope, nom_el, score) in enumerate(students, start=1):
            if previous_score is not None and score != previous_score:
                current_rank = idx
                
            cursor.execute("""
                UPDATE claselves
                set moyanuel=%s + %s WHERE id_el=%s and id_cl=%s and id_ann=%s 
                
            """, (id_el, id_cl, id_ann))
            
            previous_score = score

        conn.commit()
        flash('تم تحديث ترتيب الطلاب بنجاح', 'success')
        return True

    except psycopg2.Error as e:
        conn.rollback()
        flash(f'خطأ في قاعدة البيانات: {str(e)}', 'danger')
        return False
        
    except Exception as e:
        flash(f'حدث خطأ غير متوقع: {str(e)}', 'danger')
        return False
        
    finally:
        if 'cursor' in locals():
            cursor.close()

@app.route('/fraisscol', methods=['GET', 'POST'])
def fraisscol():
    if not session.get('logged_in'):
        return redirect(url_for('login'))    
    if 'annee_info' not in session:
        flash('الرجاء اختيار سنة دراسية أولاً', 'warning')
        return redirect(url_for('annees'))
    if 'selected_cl' not in session:
        flash('الرجاء اختيار صف دراسي أولاً', 'warning')
        return redirect(url_for('select_classefs'))
    
    id_ann = session['annee_info']['id_ann'] 
    id_ope = session.get('selected_ope')
    id_cl = session['selected_cl']   
    conn = None
    cur = None    
    try:
        conn = get_db_connection()
        cur = conn.cursor() 
        
        # استعلام المصروفات الحالية مع معلومات الطالب
        cur.execute('''
            SELECT id_cc, id_el, num_el, nom_el, credit_cfa, discription 
            FROM contenu_comtable NATURAL JOIN eleves
            WHERE id_ope = %s AND id_ann = %s
            ORDER BY id_cc DESC
        ''', (id_ope, id_ann))
        fraisscole_liste = cur.fetchall() 

        # استعلام جميع طلاب الصف (بدون استثناء)
        cur.execute('''
            SELECT id_el, num_el, nom_el FROM eleves NATURAL JOIN claselves 
            WHERE id_ann = %s AND id_cl = %s 
            ORDER BY nom_el, id_el DESC
        ''', (id_ann, id_cl))
        all_students = cur.fetchall()
        
        return render_template('fraisscol.html',
                            fraisscol_list=all_students,
                            fraisscole_liste=fraisscole_liste)
                
    except Exception as e:
        flash(f'حدث خطأ في قاعدة البيانات: {str(e)}', 'danger')
        app.logger.error(f'خطأ في قاعدة البيانات: {str(e)}')
        return redirect(url_for('select_classefs'))
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()

@app.route('/select_opeconelfc', methods=['GET', 'POST'])
def select_opeconelfc(): 
    if not session.get('logged_in'):
        return redirect(url_for('login'))    
    if 'annee_info' not in session:
        flash('الرجاء اختيار سنة دراسية أولاً', 'warning')
        return redirect(url_for('annees'))    
    id_ann = session['annee_info']['id_ann']  
    if request.method == 'POST':        
        session['selected_ope'] = request.form['id_ope']
        return redirect(url_for('select_classefs'))   
    conn = None
    cur = None    
    try:
        conn = get_db_connection()
        cur = conn.cursor() 
        cur.execute('''
        SELECT id_ope, no_ope, date_ope FROM operation_eleve 
        where id_ann=%s and id_ope !=683 ORDER BY id_ope DESC        
        ''', (id_ann,))
        operations = cur.fetchall()              
        return render_template('select_opeconelfs.html',
                              operations=operations)    
    except Exception as e:
        flash(f'حدث خطأ في قاعدة البيانات: {str(e)}', 'danger')
        app.logger.error(f'خطأ في قاعدة البيانات: {str(e)}')
        return redirect(url_for('menu'))
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()

@app.route('/select_classefs', methods=['GET', 'POST'])
def select_classefs():
    if 'selected_ope' not in session:
        return redirect(url_for('select_opeconelfc'))
    if request.method == 'POST':
        session['selected_cl'] = request.form['id_cl']
        return redirect(url_for('fraisscol'))
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute('SELECT id_cl, nom_cl FROM classes')
    classes = cur.fetchall()
    conn.close()
    return render_template('select_classe.html', classes=classes)

@app.route('/ajouter_fraisscol', methods=['POST'])
def ajouter_fraisscol():
    if not session.get('logged_in'):
        return redirect(url_for('login'))    
    # احصل على البيانات من النموذج
    id_ann = session['annee_info']['id_ann'] 
    id_ope = session['selected_ope'] 
    id_el = request.form['id_el']    
    credit_cfa = request.form['credit_cfa']
    discription = request.form['discription']    
    try:
        conn = get_db_connection()
        cur = conn.cursor()       
        # أدخل البيانات في الجدول
        cur.execute('''
            INSERT INTO contenu_comtable (id_el, id_ope, credit_cfa, discription, id_ann)
            VALUES (%s, %s, %s, %s, %s)
        ''', (id_el, id_ope, credit_cfa, discription, id_ann))

        cur.execute('''
            SELECT sum(credit_cfa) as sum_creditcfa FROM contenu_comtable 
            where id_el=%s and id_ann=%s            
        ''', (id_el, id_ann))
        sum_credit_cfa = cur.fetchone()[0] or 0  # استخدام 0 إذا كانت القيمة NULL

        cur.execute('''
            UPDATE claselves SET credit_cfaa=%s 
            WHERE id_el=%s AND id_ann=%s
        ''', (sum_credit_cfa, id_el, id_ann))
        
        conn.commit()
        flash('تمت إضافة المصروفات الدراسية بنجاح', 'success')
    except Exception as e:
        conn.rollback()
        flash(f'حدث خطأ أثناء إضافة المصروفات: {str(e)}', 'danger')
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()
    
    return redirect(url_for('fraisscol'))

@app.route('/modifier_fraisscol', methods=['POST'])
def modifier_fraisscol():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    id_ann = session['annee_info']['id_ann'] 
    id_ope = session['selected_ope']

    id_cc = request.form['id_cc']
    id_el = request.form['id_el']
    credit_cfa = request.form['credit_cfa']
    discription = request.form['discription']
    
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            UPDATE contenu_comtable 
            SET id_el=%s, credit_cfa=%s, discription=%s
            WHERE id_cc=%s
        """, (id_el, credit_cfa, discription, id_cc))

        cur.execute('''
            SELECT sum(credit_cfa) as sum_creditcfa FROM contenu_comtable 
            where id_el=%s and id_ann=%s            
        ''', (id_el, id_ann))
        sum_credit_cfa = cur.fetchone()[0] or 0  # استخدام 0 إذا كانت القيمة NULL

        cur.execute('''
            UPDATE claselves SET credit_cfaa=%s 
            WHERE id_el=%s AND id_ann=%s
        ''', (sum_credit_cfa, id_el, id_ann))
        
        conn.commit()
        
        flash('تم تعديل المصروفات الدراسية بنجاح', 'success')
    except Exception as e:
        conn.rollback()
        flash(f'حدث خطأ أثناء التعديل: {str(e)}', 'danger')
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()
    
    return redirect(url_for('fraisscol'))

#@app.route('supprimer_fraisscol/<int:id>', methods=['GET'])
#def supprimer_fraisscol(id):
#    pass
#    if not session.get('logged_in'):
#        return redirect(url_for('login'))
#    
#    id_cc = request.form['id_cc']
#     
#    try:
#        conn = get_db_connection()
#        cur = conn.cursor()
#        cur.execute("""
#            delete from contenu_comtable WHERE id_cc=%s
#        """, (id_cc))
#        conn.commit()
#        flash('تم تعديل المصروفات الدراسية بنجاح', 'success')
#    except Exception as e:
#        conn.rollback()
#        flash(f'حدث خطأ أثناء التعديل: {str(e)}', 'danger')
#    finally:
#        if cur:
#            cur.close()
#        if conn:
#            conn.close()
#    
#    return redirect(url_for('fraisscol'))
#===================================الطباعة =======================

@app.route('/opeconelfc')
def opeconelfc():
    if not session.get('logged_in'):
        return redirect(url_for('login'))    
    if 'annee_info' not in session:
        flash('الرجاء اختيار سنة دراسية أولاً', 'warning')
        return redirect(url_for('annees'))    
    id_ann = session['annee_info']['id_ann']    
    conn = None
    cur = None    
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        cur.execute("""
            SELECT id_ope, no_ope, date_ope 
            FROM operation_eleve 
            WHERE no_ope != 0 AND id_ann = %s
           """, (id_ann,))  # استخدام id_ann بدلاً من session.get('id_ann')
        
        # جلب الفصول
        cur.execute('SELECT id_cl, nom_cl FROM classes')
        classes = cur.fetchall()       
        # استعلام معدل: جلب بيانات الطلاب مع اسم الصف
        cur.execute('''
            SELECT id_el, num_el, nom_el, tel, gender, ader_parent, eta_etuditn, type_el, id_cl, nom_cl  
            FROM eleves NATURAL JOIN claselves NATURAL JOIN classes 
            WHERE id_ann = %s
        ''', (id_ann,))
        eleves_list = cur.fetchall()       
        return render_template('eleves.html',
                            classes=classes,
                            eleves_list=eleves_list)       
                
    except Exception as e:
        flash(f'حدث خطأ في قاعدة البيانات: {str(e)}', 'danger')
        app.logger.error(f'خطأ في قاعدة البيانات: {str(e)}')
        return redirect(url_for('menu'))
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()

@app.route('/opeconel', methods=['GET', 'POST'])
def opeconel():
    # التحقق من المصادقة ووجود سنة دراسية
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    if 'annee_info' not in session:
        flash('الرجاء اختيار سنة دراسية أولاً', 'warning')
        return redirect(url_for('annees'))   
    # الحصول على id_ann من session
    id_ann = session['annee_info']['id_ann']
    
    conn = None
    cur = None
    try:
        # جلب البيانات الأولية
        conn = get_db_connection()
        cur = conn.cursor()
        
        # جلب بيانات المركز التعليمي - استخدام id_ann الصحيح
        cur.execute("""
            SELECT id_ann, num_ann, annee_sc, nom_cnt, adressear, adressefr, tel, nom_cntfr, nom_cntdet, nom_cntdetfr 
            FROM ann_scolaire NATURAL JOIN nom_centere 
            WHERE id_ann = %s
        """, (id_ann,))  # استخدام id_ann بدلاً من session.get('id_ann')
        
        center_data = cur.fetchone()
        
        if not center_data:
            flash('لا توجد بيانات للمركز التعليمي لهذه السنة', 'error')
            return redirect(url_for('dashboard'))
        
        # جلب عمليات الطلاب للسنة الحالية
        cur.execute("""
            SELECT id_ope, no_ope, date_ope 
            FROM operation_eleve 
            WHERE no_ope != 0 AND id_ann = %s
            ORDER BY id_ope DESC
        """, (id_ann,))
        
        operations = cur.fetchall()
        
        return render_template('opeconel.html', 
                            center_data=center_data,
                            operations=operations,
                            id_ann=id_ann)  # تمرير id_ann إلى القالب
        
    except Exception as e:
        app.logger.error(f"Error in opeconel: {str(e)}")
        flash('حدث خطأ أثناء جلب البيانات', 'error')
        return redirect(url_for('dashboard'))
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()


@app.route('/get_classes', methods=['POST'])
def get_classes():
    conn = get_db_connection()
    cur = conn.cursor()
    
    cur.execute("SELECT id_cl, nom_cl FROM classes ORDER BY id_cl DESC")
    classes = cur.fetchall()
    
    cur.close()
    conn.close()    
    return jsonify([{'id': c[0], 'name': c[1]} for c in classes])

@app.route('/get_students', methods=['POST'])
def get_students():
    data = request.json
    id_ann = data.get('id_ann')
    id_cl = data.get('id_cl')
    id_ope = data.get('id_ope')
    
    conn = get_db_connection()
    cur = conn.cursor()
    
    cur.execute("""
        SELECT id_el, num_el, nom_el 
        FROM eleves NATURAL JOIN claselves NATURAL JOIN classes 
        WHERE id_ann = %s AND id_cl = %s AND id_el NOT IN (
            SELECT id_el FROM contenu_comtable WHERE id_ope = %s
        ) 
        ORDER BY nom_el, id_el DESC
    """, (id_ann, id_cl, id_ope))
    
    students = cur.fetchall()
    cur.close()
    conn.close()    
    return jsonify([{'id': s[0], 'num': s[1], 'name': s[2]} for s in students])

@app.route('/add_operation', methods=['POST'])
def add_operation():
    id_ann = data.get('id_ann')
    id_cl = data.get('id_cl')
    id_ope = data.get('id_ope')
    con = None
    cur = None
    try:
        data = request.form
        
        # التحقق من البيانات المطلوبة
        required_fields = ['credit_cfa', 'id_el', 'id_ope', 'id_ann']
        for field in required_fields:
            if not data.get(field):
                return jsonify({'success': False, 'message': f'حقل {field} مطلوب'}), 400

        con = psycopg2.connect(**DB_CONFIG)
        cur = con.cursor()

        # 1. التحقق من وجود السجل مع قفل الصف
        cur.execute("""
            SELECT idclel, credit_cfaa 
            FROM claselves 
            WHERE id_el = %s AND id_ann = %s
            FOR UPDATE
        """, (int(data['id_el']), int(data['id_ann'])))
        
        existing_record = cur.fetchone()
        
        if not existing_record:
            return jsonify({
                'success': False,
                'message': 'السجل غير موجود في claselves',
                'details': {
                    'id_el': data['id_el'],
                    'id_ann': data['id_ann'],
                    'table': 'claselves'
                }
            }), 404

        # 2. إضافة العملية الجديدة
        cur.execute("""
            INSERT INTO contenu_comtable 
            (credit_cfa, id_el, id_ope, id_ann, discription) 
            VALUES (%s, %s, %s, %s, %s)
            RETURNING id_cc
        """, (
            float(data['credit_cfa']),
            int(data['id_el']),
            int(data['id_ope']),
            int(data['id_ann']),
            data.get('discription', '')
        ))
        inserted_id = cur.fetchone()[0]

        # 3. حساب الرصيد الجديد (المجموع التراكمي)
        cur.execute("""
            SELECT COALESCE(SUM(credit_cfa), 0) 
            FROM contenu_comtable 
            WHERE id_el = %s AND id_ann = %s
        """, (int(data['id_el']), int(data['id_ann'])))
        sum_credit = float(cur.fetchone()[0])

        # 4. التحديث باستخدام idclel كمفتاح رئيسي
        cur.execute("""
            UPDATE claselves 
            SET credit_cfaa = %s 
            WHERE id_el = %s and id_ann=%s
            RETURNING credit_cfaa
        """, (sum_credit, inserted_id, id_ann))
        
        updated_record = cur.fetchone()
        
        if not updated_record:
            raise Exception("فشل التحديث: لم يتم العثور على السجل بعد التحديث")

        print(f"تم التحديث بنجاح: الرصيد الجديد {updated_record[0]}")

        con.commit()
        return jsonify({
            'success': True,
            'message': 'تمت العملية بنجاح',
            'data': {
                'inserted_id': inserted_id,
                'old_balance': existing_record[1],
                'new_balance': updated_record[0],
                'idclel': existing_record[0]
            }
        })

    except psycopg2.Error as e:
        if con:
            con.rollback()
        error_msg = f"خطأ في قاعدة البيانات: {e.pgerror}" if e.pgerror else str(e)
        print(error_msg)
        return jsonify({
            'success': False,
            'message': 'خطأ في قاعدة البيانات',
            'error': error_msg
        }), 500
        
    finally:
        if cur:
            cur.close()
        if con:
            con.close()
            
@app.route('/print_receipt', methods=['POST'])
def print_receipt():
    data = request.json
    try:
        # إنشاء ملف الوورد
        doc = DocxTemplate("templete_Recu/recu_eleves.docx")
        doc.render(data)
        output_path = os.path.abspath('Imprimer_Eleves/Imprim_Recu_eleve.docx')
        doc.save(output_path)
        
        return send_file(output_path, as_attachment=True)
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})
    
#------------------EXCEL---------------------------
@app.route('/excelindex')
def excelindex():
    if not session.get('logged_in'):
        return redirect(url_for('login'))    
    if 'annee_info' not in session:
        flash('الرجاء اختيار سنة دراسية أولاً', 'warning')
        return redirect(url_for('annees'))    
    id_ann = session['annee_info']['id_ann']    
    conn = None
    cur = None    
    try:
        conn = get_db_connection()
        cur = conn.cursor() 
        # جلب الفصول
        cur.execute('SELECT id_cl, nom_cl FROM classes')
        classes = cur.fetchall()      
        return render_template('excelindex.html',
                            classes=classes)       
                
    except Exception as e:
        flash(f'حدث خطأ في قاعدة البيانات: {str(e)}', 'danger')
        app.logger.error(f'خطأ في قاعدة البيانات: {str(e)}')
        return redirect(url_for('menu'))
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()


@app.route('/get_classess')
def get_classess():
    if not session.get('logged_in'):
        return redirect(url_for('login'))    
    if 'annee_info' not in session:
        flash('الرجاء اختيار سنة دراسية أولاً', 'warning')
        return redirect(url_for('annees'))    
    id_ann = session['annee_info']['id_ann']    
    conn = None
    cur = None    
    try:
        conn = get_db_connection()
        cur = conn.cursor() 
        # جلب الفصول
        cur.execute('SELECT id_cl, nom_cl FROM classes')
        classes = cur.fetchall()      
        return render_template('excelindex.html',
                            classes=classes)       
                
    except Exception as e:
        flash(f'حدث خطأ في قاعدة البيانات: {str(e)}', 'danger')
        app.logger.error(f'خطأ في قاعدة البيانات: {str(e)}')
        return redirect(url_for('menu'))
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()

@app.route('/import_from_excel', methods=['POST'])  
def import_from_excel():
    if not session.get('logged_in'):
        return redirect(url_for('login'))    
    if 'annee_info' not in session:
        flash('الرجاء اختيار سنة دراسية أولاً', 'warning')
        return redirect(url_for('annees'))
    
    if 'file' not in request.files:
        flash('لم يتم اختيار ملف', 'error')
        return redirect(request.url)
    
    file = request.files['file']
    if file.filename == '':
        flash('لم يتم اختيار ملف', 'error')
        return redirect(request.url)
    
    if not file.filename.endswith('.xlsx'):
        flash('الملف يجب أن يكون بصيغة Excel (.xlsx)', 'error')
        return redirect(request.url)
    
    try:
        id_ann = session['annee_info']['id_ann']
        df = pd.read_excel(file)
        
        con = psycopg2.connect(**DB_CONFIG)
        cur = con.cursor()
        
        for row in df.itertuples():
            cur.execute(
                "INSERT INTO excel (num_el, nom_el, tel, gender, ader_parent, montant_arier, inscrpt, eta_etuditn, montantp) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)",
                (
                    row.num_el,
                    row.nom_el,
                    row.tel,
                    row.gender,
                    row.ader_parent,
                    row.montant_arier,
                    row.inscrpt,
                    row.eta_etuditn,
                    row.montantp
                )
            )
        
        con.commit()
        flash('تمت الاستيراد بنجاح', 'success')
        return redirect(url_for('excelindex'))
    except Exception as e:
        flash(f'خطأ في الاستيراد: {str(e)}', 'error')
        return redirect(request.url)
    finally:
        if 'cur' in locals(): cur.close()
        if 'con' in locals(): con.close()

@app.route('/add_to_database', methods=['POST'])        
def add_to_database():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    if 'annee_info' not in session:
        flash('الرجاء اختيار سنة دراسية أولاً', 'warning')
        return redirect(url_for('annees'))
    
    try:
        con = psycopg2.connect(**DB_CONFIG)
        cur = con.cursor()     
        id_ann = session['annee_info']['id_ann'] 
        id_cl = request.form.get('id_cl')
        
        if not id_cl:
            flash('الرجاء اختيار الفصل أولاً', 'error')
            return redirect(url_for('excelindex'))
        
        # جلب id_ope كقيمة مفردة
        cur.execute("SELECT id_ope FROM operation_eleve WHERE id_ope=683 LIMIT 1")
        result = cur.fetchone()
        id_ope = result[0] if result else 1  # قيمة افتراضية إذا لم تكن موجودة

        # جلب مصاريف الفصل
        cur.execute("""
            SELECT type_frais, montant_a_paiye, inscrption 
            FROM classes 
            NATURAL JOIN clas_frais_scol 
            NATURAL JOIN frais_scol 
            WHERE id_cl=%s 
            ORDER BY montant_a_paiye
        """, (id_cl,))
        fees = cur.fetchall()      
        # جلب الطلاب من جدول excel
        cur.execute("SELECT num_el, nom_el, tel, gender, ader_parent, montant_arier, inscrpt, eta_etuditn, montantp FROM excel")
        excel_students = cur.fetchall()
        
        if not excel_students:
            flash('لا يوجد طلاب لاضافتهم', 'warning')
            return redirect(url_for('excelindex'))
        
        # إدخال في الجداول الرئيسية
        for student in excel_students:
            # إدخال في جدول الطلاب
            cur.execute("""
                INSERT INTO eleves (num_el, nom_el, tel, gender, ader_parent, mon_exl, inscrpt, eta_etudition, montantpey)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id_el
            """, (
                student[0], student[1], student[2], student[3], 
                student[4], student[5], student[6], student[7], student[8]
            ))
            student_id = cur.fetchone()[0]            
            
            # تحديد المصاريف
            montant_a_paiye = 0
            for fee in fees:
                if student[7] == fee[0]:
                    montant_a_paiye = fee[1]
                    break
                
            # إدخال في claselves
            if student[6] > 0:  # إذا كان مسجلاً
                cur.execute("""
                    INSERT INTO claselves 
                    (type_el, id_el, eta_etuditn, montant_arier, inscrption, id_ann, id_cl, montant_a_paiye, credit_cfaa) 
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    'Nouveau', student_id, student[7], student[5], 
                    student[6], id_ann, id_cl, montant_a_paiye, student[8]
                ))
            else:
                cur.execute("""
                    INSERT INTO claselves 
                    (id_el, eta_etuditn, montant_arier, id_ann, id_cl, montant_a_paiye, credit_cfaa) 
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (
                    student_id, student[7], student[5], 
                    id_ann, id_cl, montant_a_paiye, student[8]
                ))            
            
            # إدخال في المحاسبة
            cur.execute("""
                INSERT INTO contenu_comtable 
                (id_el, id_ope, id_ann, credit_cfa) 
                VALUES (%s, %s, %s, %s)
            """, (student_id, id_ope, id_ann, student[8]))
        
        # حذف بيانات excel بعد الانتهاء
        cur.execute('TRUNCATE TABLE excel')
        con.commit()
        
        flash('تمت إضافة الطلاب إلى قاعدة البيانات بنجاح', 'success')
        return redirect(url_for('excelindex'))
        
    except Exception as e:
        con.rollback()
        flash(f'حدث خطأ أثناء الإضافة: {str(e)}', 'error')
        app.logger.error(f'خطأ في إضافة الطلاب: {str(e)}')
        return redirect(url_for('excelindex'))
    finally:
        if 'cur' in locals(): cur.close()
        if 'con' in locals(): con.close()
#==========================================PRINT=========================================
#---b-------طباعة فاتورة الطالب مع generate_student_invoice + fraisscol.html----------------------
@app.route('/generate_student_invoice', methods=['POST'])
def generate_student_invoice():
    try:
        data = request.json
        student_id = data.get('student_id')
        #id_ope = data.get('id_ope')
        id_ope = session['selected_ope']
        id_ann = data.get('id_ann')        
        if not student_id or not id_ann:
            return jsonify({"status": "error", "message": "معرّف الطالب أو السنة الدراسية مفقود"}), 400
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        # الحصول على معلومات المركز من قاعدة البيانات
        cur.execute("""
            SELECT id_ann, num_ann, annee_sc, nom_cnt, adressear, adressefr, 
            tel, nom_cntfr, nom_cntdet, nom_cntdetfr 
            FROM ann_scolaire NATURAL JOIN nom_centere 
            WHERE id_ann = %s 
        """, (id_ann,))
        ecol_info = cur.fetchone()
        if not ecol_info:
            return jsonify({"status": "error", "message": "لم يتم العثور على معلومات المركز"}), 404
        # استخراج بيانات المركز
        nom_cnt = ecol_info[3]
        adressear = ecol_info[4]
        adressefr = ecol_info[5]
        tel = ecol_info[6]
        nom_cntfr = ecol_info[7]
        nom_cntdet = ecol_info[8]
        nom_cntdetfr = ecol_info[9]
        # 1. الحصول على بيانات الطالب الأساسية
        cur.execute("""
            SELECT id_el, nom_el, nom_cl, annee_sc 
            FROM eleves 
            NATURAL JOIN claselves 
            NATURAL JOIN classes 
            NATURAL JOIN ann_scolaire 
            WHERE id_el = %s AND id_ann = %s
        """, (student_id, id_ann))
        student_info = cur.fetchone()
      
        if not student_info:
            return jsonify({"status": "error", "message": "لم يتم العثور على الطالب"}), 404
        id_el, nom_el, nom_cl, annee_scol = student_info
           # 3. الحصول على بيانات الدفعات
        cur.execute("""
            SELECT no_ope, date_ope, credit_cfa, discription from operation_eleve natural join contenu_comtable natural join eleves 
            WHERE id_el = %s AND id_ope = %s
        """, (student_id, id_ope))
        detail_list2 = cur.fetchall()

        cur.execute("""
            SELECT distinct montant_a_paiye, inscrption, montant_arier, (montant_a_paiye+inscrption+montant_arier) as montant_total_apaiye, 
            credit_cfaa as som_total_payee, (montant_a_paiye+inscrption+montant_arier)-credit_cfaa as solde 
            FROM eleves NATURAL JOIN contenu_comtable NATURAL JOIN classes NATURAL JOIN claselves NATURAL JOIN ann_scolaire 
            WHERE id_el = %s AND id_ann = %s 
        """, (student_id, id_ann))
        total_paiment = cur.fetchall()
        # 4. حساب المجاميع
        montantapaiye = sum(item[0] for item in total_paiment) if total_paiment else 0
        inscrptions = sum(item[1] for item in total_paiment) if total_paiment else 0
        montantarier = sum(item[2] for item in total_paiment) if total_paiment else 0
        sod = sum(item[3] for item in total_paiment) if total_paiment else 0
        mottpayer = sum(item[4] for item in total_paiment) if total_paiment else 0
        resmont = sum(item[5] for item in total_paiment) if total_paiment else 0
        # 5. تعريف context بشكل صحيح
        context = {
            "nomcnt": nom_cnt,
            "adressear": adressear,
            "adressefr": adressefr,
            "tel": tel,
            "nomcntfr": nom_cntfr,
            "nomcntdet": nom_cntdet,
            "nomcntdetfr": nom_cntdetfr,
            "nomel": nom_el,
            "nomcl": nom_cl,
            "anneesc": annee_scol,
            "total_paiment": total_paiment,
            "detail_list2": detail_list2,
            "montantapaiye": montantapaiye,
            "inscrptions": inscrptions,
            "montantarier": montantarier,
            "sod": sod,
            "mottpayer": mottpayer,
            "resmont": resmont
        }
        BASE_DIR = Path(__file__).parent
        template_path = BASE_DIR / "templete_Recu"/"recu_eleves.docx"
        output_dir = BASE_DIR / "Imprimer_Eleves"        
        os.makedirs(output_dir, exist_ok=True)        
        if not template_path.exists():
            return jsonify({
                "status": "error",
                "message": f"ملف القالب غير موجود في: {template_path}"
            }), 404

        # إنشاء وحفظ الوثيقة
        doc = DocxTemplate(str(template_path))
        doc.render(context)
        output_filename = f"recu_eleves_{student_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = output_dir / output_filename
        doc.save(str(output_path))        
        # إرجاع استجابة JSON فقط
        return jsonify({
            "status": "success",
            
            "message": "تم حفظ الملف بنجاح على الخادم",
            "file_info": {
                "name": output_filename,
                "path": str(output_path),
                "student_name": nom_el,
                "class": nom_cl,
                "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
        }), 200

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        if 'conn' in locals():
            conn.close()     
#--------------------------تفاصيل مدفوعات الطالب  generate_student_payments+ fraisscol.html
@app.route('/generate_student_payments', methods=['POST'])
def generate_student_payments():
    try:
        data = request.json
        student_id = data.get('student_id')
        id_ann = data.get('id_ann')        
        if not student_id or not id_ann:
            return jsonify({"status": "error", "message": "معرّف الطالب أو السنة الدراسية مفقود"}), 400
        
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        # الحصول على تفاصيل المدفوعات
        cur.execute("""
            SELECT no_ope, date_ope, credit_cfa, discription 
            FROM operation_eleve 
            NATURAL JOIN contenu_comtable 
            WHERE id_el = %s AND id_ann = %s AND no_ope != 0
            ORDER BY date_ope
        """, (student_id, id_ann))
        payments = cur.fetchall()

        # معلومات المركز
        cur.execute("""
            SELECT nom_cnt, adressear, adressefr, tel, 
                   nom_cntfr, nom_cntdet, nom_cntdetfr 
            FROM ann_scolaire 
            NATURAL JOIN nom_centere 
            WHERE id_ann = %s 
        """, (id_ann,))
        ecol_info = cur.fetchone()
        
        if not ecol_info:
            return jsonify({"status": "error", "message": "لم يتم العثور على معلومات المركز"}), 404

        nom_cnt, adressear, adressefr, tel, nom_cntfr, nom_cntdet, nom_cntdetfr = ecol_info

        # بيانات الطالب
        cur.execute("""
            SELECT nom_el, nom_cl, annee_sc 
            FROM eleves 
            NATURAL JOIN claselves 
            NATURAL JOIN classes 
            NATURAL JOIN ann_scolaire 
            WHERE id_el = %s AND id_ann = %s
        """, (student_id, id_ann))
        student_info = cur.fetchone()
        
        if not student_info:
            return jsonify({"status": "error", "message": "لم يتم العثور على الطالب"}), 404

        nom_el, nom_cl, annee_scol = student_info

        # بيانات الدفعات
        cur.execute("""
            SELECT DISTINCT montant_a_paiye, inscrption, montant_arier, 
                   (montant_a_paiye+inscrption+montant_arier) as montant_total,
                   credit_cfaa as total_paye,
                   (montant_a_paiye+inscrption+montant_arier) - credit_cfaa as reste
            FROM eleves 
            NATURAL JOIN contenu_comtable 
            NATURAL JOIN classes 
            NATURAL JOIN claselves 
            NATURAL JOIN ann_scolaire 
            WHERE id_el = %s AND id_ann = %s
        """, (student_id, id_ann))
        total_payments = cur.fetchall()

        if not total_payments:
            return jsonify({"status": "error", "message": "لا توجد مدفوعات مسجلة للطالب"}), 404

        # إعداد البيانات للقالب
        context = {        
            "nomcnt": nom_cnt,
            "adressear": adressear,
            "adressefr": adressefr,
            "tel": tel,
            "nomcntfr": nom_cntfr,
            "nomcntdet": nom_cntdet,
            "nomcntdetfr": nom_cntdetfr,
            "nomel": nom_el,
            "nomcl": nom_cl,
            "anneesc": annee_scol,
            "total_payments": total_payments,
            "payments": payments 
        }

        # مسارات الملفات
        BASE_DIR = Path(__file__).parent
        template_path = BASE_DIR / "templete_Recu" / "recu_detaileleves.docx"
        output_dir = BASE_DIR / "Imprimer_Eleves"
        
        os.makedirs(output_dir, exist_ok=True)
        
        if not template_path.exists():
            return jsonify({
                "status": "error",
                "message": f"ملف القالب غير موجود في: {template_path}"
            }), 404

        # إنشاء وحفظ الوثيقة
        doc = DocxTemplate(str(template_path))
        doc.render(context)
        output_filename = f"recu_detaileleves_{student_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = output_dir / output_filename
        doc.save(str(output_path))
        
        # إرجاع استجابة JSON فقط
        return jsonify({
            "status": "success",
            "message": "تم حفظ الملف بنجاح على الخادم",
            "file_info": {
                "name": output_filename,
                "path": str(output_path),
                "student_name": nom_el,
                "class": nom_cl,
                "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
        }), 200

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        if 'conn' in locals():
            conn.close()

# إذا أردت إضافة إمكانية التحميل لاحقاً
#@app.route('/download_payment/<filename>', methods=['GET'])
#def download_payment(filename):
#    try:
#        output_dir = Path(__file__).parent / "Imprimer_Eleves"
#        file_path = output_dir / filename
#        
#        if not file_path.exists():
#            return jsonify({"status": "error", "message": "الملف غير موجود"}), 404
#        
#        return send_file(
#            file_path,
#            as_attachment=True,
#            download_name=filename,
#            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#        )
#    except Exception as e:
#        return jsonify({"status": "error", "message": str(e)}), 500
#    
                                 #-------------------
def _add_subtotal_row(table, subtotal):
    """دالة مساعدة لإضافة صف المجموع الفرعي"""
    total_frais, total_inscription, total_arier, total_a_payer, total_payee, total_solde = subtotal
    row_cells = table.add_row().cells
    row_cells[0].text = "المجموع الفرعي"
    row_cells[1].text = ""
    row_cells[2].text = str(total_frais)
    row_cells[3].text = str(total_inscription)
    row_cells[4].text = str(total_arier)
    row_cells[5].text = str(total_a_payer)
    row_cells[6].text = str(total_payee)
    row_cells[7].text = str(total_solde)
    
    # تنسيق صف المجموع الفرعي
    for cell in row_cells:
        shading = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
        cell._element.get_or_add_tcPr().append(shading)

@app.route('/generate_partial_payment_report', methods=['POST'])
def generate_partial_payment_report():
    try:
        if 'annee_info' not in session:
            return jsonify({'error': 'لم يتم اختيار سنة دراسية'}), 400
        idann = session['annee_info']['id_ann']
        noman = session['annee_info']['annee_sc']
        nomcnt = session['annee_info']['nom_cnt']
        conn = psycopg2.connect(**DB_CONFIG)
        cursor = conn.cursor()
        # استعلام الطلاب الذين دفعوا جزئياً
        cursor.execute("""
            SELECT 
                nom_cl, num_el, nom_el, montant_a_paiye, inscrption, montant_arier, 
                (montant_a_paiye + inscrption + montant_arier) AS total_a_payer, 
                credit_cfaa AS total_payee, 
                (montant_a_paiye + inscrption + montant_arier) - credit_cfaa AS solde 
            FROM eleves 
            NATURAL JOIN classes 
            NATURAL JOIN claselves 
            NATURAL JOIN ann_scolaire 
            WHERE id_ann = %s 
            AND credit_cfaa < (montant_a_paiye + inscrption + montant_arier) 
            AND credit_cfaa > 0
            ORDER BY nom_cl, num_el
        """, (idann,))
        students = cursor.fetchall()
        # استعلام المجاميع الفرعية لكل فرع
        cursor.execute("""
            SELECT 
                nom_cl, 
                SUM(montant_a_paiye) AS total_frais, 
                SUM(inscrption) AS total_inscription, 
                SUM(montant_arier) AS total_arier, 
                SUM(montant_a_paiye + inscrption + montant_arier) AS total_a_payer, 
                SUM(credit_cfaa) AS total_payee, 
                SUM((montant_a_paiye + inscrption + montant_arier) - credit_cfaa) AS total_solde
            FROM eleves 
            NATURAL JOIN classes 
            NATURAL JOIN claselves 
            NATURAL JOIN ann_scolaire 
            WHERE id_ann = %s 
            AND credit_cfaa < (montant_a_paiye + inscrption + montant_arier) 
            AND credit_cfaa > 0
            GROUP BY nom_cl 
            ORDER BY nom_cl
        """, (idann,))
        subtotals = {row[0]: row[1:] for row in cursor.fetchall()}
        # استعلام المجموع الكلي
        cursor.execute("""
            SELECT 
                SUM(montant_a_paiye), 
                SUM(inscrption), 
                SUM(montant_arier), 
                SUM(montant_a_paiye + inscrption + montant_arier), 
                SUM(credit_cfaa), 
                SUM((montant_a_paiye + inscrption + montant_arier) - credit_cfaa)
            FROM eleves 
            NATURAL JOIN classes 
            NATURAL JOIN claselves 
            NATURAL JOIN ann_scolaire 
            WHERE id_ann = %s 
            AND credit_cfaa < (montant_a_paiye + inscrption + montant_arier) 
            AND credit_cfaa > 0
        """, (idann,))
        grand_total = cursor.fetchone()
        # إنشاء مستند Word
        doc = Document()
        for section in doc.sections:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
        # إضافة معلومات التقرير
        current_date = datetime.datetime.now().strftime('%Y-%m-%d')
        doc.add_paragraph(f"تاريخ الطباعة: {current_date}")
        doc.add_paragraph(f"العام الدراسي: {noman}")
        doc.add_paragraph(f"المدرسة: {nomcnt}")
        doc.add_heading('قائمة الطلاب الذين لم يدفعوا كاملاً', 0)
        # إضافة بيانات الطلاب حسب الفروع
        current_branch = None
        for student in students:
            nom_cl, num_el, nom_el, frais, inscription, arier, total, payee, solde = student
            if nom_cl != current_branch:
                if current_branch is not None:
                    # إضافة المجموع الفرعي للفرع السابق
                    _add_subtotal_row(table, subtotals[current_branch])
                # بدء جدول جديد للفرع الحالي
                doc.add_heading(f'الفرع: {nom_cl}', level=2)
                table = doc.add_table(rows=1, cols=8)
                table.style = 'Table Grid'                
                # تحديد عرض الأعمدة
                widths = [Inches(0.8), Inches(2), Inches(1), Inches(1), Inches(1.5), Inches(1), Inches(1), Inches(1)]
                for i, width in enumerate(widths):
                    table.columns[i].width = width                
                # إضافة رؤوس الأعمدة
                headers = ['الرقم', 'اسم الطالب', 'الرسوم', 'التسجيل', 'المتأخرات', 'المجموع', 'المدفوع', 'المتبقي']
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header                
                current_branch = nom_cl
            # إضافة بيانات الطالب
            row_cells = table.add_row().cells
            row_cells[0].text = num_el
            row_cells[1].text = nom_el
            row_cells[2].text = str(frais)
            row_cells[3].text = str(inscription)
            row_cells[4].text = str(arier)
            row_cells[5].text = str(total)
            row_cells[6].text = str(payee)
            row_cells[7].text = str(solde)
        # إضافة المجموع الفرعي للفرع الأخير
        if current_branch is not None and current_branch in subtotals:
            _add_subtotal_row(table, subtotals[current_branch])
        # إضافة المجموع الكلي
        doc.add_heading('المجموع الكلي لجميع الفروع', level=2)
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'        
        # تحديد عرض الأعمدة
        widths = [Inches(0.8), Inches(2), Inches(1), Inches(1), Inches(1.5), Inches(1), Inches(1), Inches(1)]
        for i, width in enumerate(widths):
            table.columns[i].width = width        
        # إضافة رؤوس الأعمدة
        headers = ['', '', 'الرسوم', 'التسجيل', 'المتأخرات', 'المجموع', 'المدفوع', 'المتبقي']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header        
        # إضافة صف المجموع الكلي
        row_cells = table.add_row().cells
        row_cells[0].text = "المجموع الكلي"
        row_cells[1].text = ""
        row_cells[2].text = str(grand_total[0])
        row_cells[3].text = str(grand_total[1])
        row_cells[4].text = str(grand_total[2])
        row_cells[5].text = str(grand_total[3])
        row_cells[6].text = str(grand_total[4])
        row_cells[7].text = str(grand_total[5])
        
        # تنسيق صف المجموع الكلي
        for cell in row_cells:
            shading = parse_xml(r'<w:shd {} w:fill="FFD966"/>'.format(nsdecls('w')))
            cell._element.get_or_add_tcPr().append(shading)

        # حفظ التقرير مؤقتاً
        if not os.path.exists('temp'):
            os.makedirs('temp')
        
        file_path = os.path.join('temp', 'قائمة الطلاب الذين لم يدفعوا كاملاً.docx')
        doc.save(file_path)
        
        # إرسال الملف كاستجابة
        return send_file(
            file_path,
            as_attachment=True,
            download_name='قائمة الطلاب الذين لم يدفعوا كاملاً.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        app.logger.error(f"Error generating partial payment report: {str(e)}")
        return jsonify({'error': str(e)}), 500
    finally:
        if 'cursor' in locals():
            cursor.close()
        if 'conn' in locals():
            conn.close()

@app.route('/generate_full_payment_report2', methods=['POST'])
def generate_full_payment_report2():
    try:
        if 'annee_info' not in session:
            return jsonify({'error': 'لم يتم اختيار سنة دراسية'}), 400
        idann = session['annee_info']['id_ann']
        noman = session['annee_info']['annee_sc']
        nomcnt = session['annee_info']['nom_cnt']

        conn = psycopg2.connect(**DB_CONFIG)
        cursor = conn.cursor()
        # استعلام الطلاب الذين دفعوا كاملاً
        cursor.execute("""
            SELECT 
                nom_cl, num_el, nom_el, montant_a_paiye, inscrption, montant_arier, 
                (montant_a_paiye + inscrption + montant_arier) AS total_a_payer, 
                credit_cfaa AS total_payee, 
                (montant_a_paiye + inscrption + montant_arier) - credit_cfaa AS solde 
            FROM eleves NATURAL JOIN classes NATURAL JOIN claselves NATURAL JOIN ann_scolaire 
            WHERE id_ann = %s 
            AND credit_cfaa >= (montant_a_paiye + inscrption + montant_arier)
            ORDER BY nom_cl, num_el
        """, (idann,))
        students1 = cursor.fetchall()
        # إنشاء مستند Word
        doc = Document()
        for section in doc.sections:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
        # إضافة معلومات التقرير
        current_date = datetime.datetime.now().strftime('%Y-%m-%d')
        doc.add_paragraph(f"تاريخ الطباعة: {current_date}")
        doc.add_paragraph(f"العام الدراسي: {noman}")
        doc.add_paragraph(f"المدرسة: {nomcnt}")
        doc.add_heading('قائمة الطلاب الذين دفعوا كامل الرسوم', 0)
        # إضافة جدول البيانات
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'       
        # رؤوس الأعمدة
        headers = ['الفرع', 'رقم الطالب', 'اسم الطالب', 'الرسوم', 'التسجيل', 'المتأخرات', 'المجموع', 'المدفوع']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        # إضافة بيانات الطلاب
        for student in students1:
            row_cells = table.add_row().cells
            for i, value in enumerate(student[:8]):  # نأخذ أول 8 قيم (لا نعرض الرصيد)
                row_cells[i].text = str(value)
        # حفظ الملف مؤقتاً
        file_path = os.path.join('temp', 'قائمة الطلاب الذين دفعوا كاملا.docx')
        doc.save(file_path)        
        return send_file(
            file_path,
            as_attachment=True,
            download_name='قائمة_الطلاب_الذين_دفعوا_كاملاً.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        app.logger.error(f"Error generating full payment report: {str(e)}")
        return jsonify({'error': str(e)}), 500
    finally:
        if 'cursor' in locals():
            cursor.close()
        if 'conn' in locals():
            conn.close()

@app.route('/generate_nopayment_report', methods=['POST'])
def generate_nopayment_report():
    conn = None
    cursor = None
    try:
        if 'annee_info' not in session:
            return jsonify({'error': 'لم يتم اختيار سنة دراسية'}), 400
        idann = session['annee_info']['id_ann']
        noman = session['annee_info']['annee_sc']
        nomcnt = session['annee_info']['nom_cnt']    
        # إنشاء اتصال بقاعدة البيانات
        conn = psycopg2.connect(**DB_CONFIG)
        cursor = conn.cursor()
        cursor.execute("""
            select nom_cl, num_el, nom_el, montant_a_paiye, inscrption, montant_arier, (montant_a_paiye+inscrption+montant_arier) 
            as Totat_a_payer, credit_cfaa as total_payee, (montant_a_paiye+inscrption+montant_arier)-credit_cfaa as solde from eleves 
            natural join classes natural join claselves natural join ann_scolaire
            WHERE id_ann = %s 
            AND credit_cfaa = 0 
        """, (idann,))
        students = cursor.fetchall()
        # استعلام المجاميع الفرعية لكل فرع
        cursor.execute("""
            select nom_cl, sum(montant_a_paiye) AS total_frais, sum(inscrption) AS total_inscription, sum(montant_arier) AS to_montantarier, 
            sum(montant_a_paiye+inscrption+montant_arier) AS to_montantapaye, sum(credit_cfaa) as total_montantpaye, 
            sum(montant_a_paiye)+sum(inscrption)+sum(montant_arier)-sum(credit_cfaa) as total_reste  from eleves natural join classes 
            natural join claselves natural join ann_scolaire
            WHERE id_ann = %s and credit_cfaa =0 group by nom_cl
        """, (idann,))        
        subtotals = cursor.fetchall()          
        # استعلام المجموع الكلي
        cursor.execute("""
            SELECT sum(montant_a_paiye), sum(inscrption), sum(montant_arier), sum(montant_a_paiye)+sum(inscrption)+sum(montant_arier), 
            sum(credit_cfaa), sum(montant_a_paiye)+sum(inscrption)+sum(montant_arier)-sum(credit_cfaa) from eleves natural join classes 
            natural join claselves natural join ann_scolaire
            WHERE id_ann = %s
            AND credit_cfaa = 0
        """, (idann,))
        grand_total = cursor.fetchone()
        # إنشاء مستند Word مع تنسيق محسّن
        doc = Document()        
        # إعداد صفحة العرض (Landscape)
        for section in doc.sections:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # إضافة معلومات التقرير
        title = doc.add_paragraph('قائمة الطلاب الذين لم يدفعوا شيئاً')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(14)
        title.runs[0].font.bold = True

        doc.add_paragraph(f"المدرسة: {nomcnt}")
        doc.add_paragraph(f"العام الدراسي: {noman}")
        doc.add_paragraph(f"تاريخ التقرير: {datetime.datetime.now().strftime('%Y-%m-%d')}")

        # إضافة بيانات الطلاب حسب الفروع
        current_branch = None
        for student in students:
            nom_cl, num_el, nom_el, frais, inscription, arier, total, payee, solde = student

            if nom_cl != current_branch:
                if current_branch is not None:
                    # إضافة المجموع الفرعي للفرع السابق
                    for subtotal in subtotals:
                        if subtotal[0] == current_branch:
                            table.add_row()
                            last_row = table.rows[-1].cells
                            last_row[0].text = "المجموع"
                            last_row[1].text = ""
                            last_row[2].text = f"{subtotal[1]:,.2f}"
                            last_row[3].text = f"{subtotal[2]:,.2f}"
                            last_row[4].text = f"{subtotal[3]:,.2f}"
                            last_row[5].text = f"{subtotal[4]:,.2f}"
                            last_row[6].text = f"{subtotal[5]:,.2f}"
                            last_row[7].text = f"{subtotal[6]:,.2f}"
                            break

                # بدء جدول جديد للفرع الحالي
                doc.add_paragraph(f"الفرع: {nom_cl}", style='Heading 2')
                table = doc.add_table(rows=1, cols=8)
                table.style = 'Table Grid'
                
                # تحديد عرض الأعمدة
                widths = [Inches(0.8), Inches(1.5), Inches(2), Inches(1), Inches(1), Inches(1), Inches(1), Inches(1)]
                for i, width in enumerate(widths):
                    table.columns[i].width = width
                
                # رؤوس الأعمدة
                headers = ['الرقم', 'الاسم', 'الرسوم', 'التسجيل', 'المتأخرات', 'المستحق', 'المدفوع', 'المتبقي']
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                current_branch = nom_cl

            # إضافة بيانات الطالب
            row_cells = table.add_row().cells
            row_cells[0].text = str(num_el)
            row_cells[1].text = nom_el
            row_cells[2].text = f"{frais:,.2f}"
            row_cells[3].text = f"{inscription:,.2f}"
            row_cells[4].text = f"{arier:,.2f}"
            row_cells[5].text = f"{total:,.2f}"
            row_cells[6].text = f"{payee:,.2f}"
            row_cells[7].text = f"{solde:,.2f}"

        # إضافة المجموع الفرعي للفرع الأخير
        if current_branch is not None:
            for subtotal in subtotals:
                if subtotal[0] == current_branch:
                    table.add_row()
                    last_row = table.rows[-1].cells
                    last_row[0].text = "المجموع"
                    last_row[1].text = ""
                    last_row[2].text = f"{subtotal[1]:,.2f}"
                    last_row[3].text = f"{subtotal[2]:,.2f}"
                    last_row[4].text = f"{subtotal[3]:,.2f}"
                    last_row[5].text = f"{subtotal[4]:,.2f}"
                    last_row[6].text = f"{subtotal[5]:,.2f}"
                    last_row[7].text = f"{subtotal[6]:,.2f}"
                    break

        # إضافة المجموع الكلي
        if grand_total and any(grand_total):
            doc.add_paragraph("المجموع الكلي", style='Heading 2')
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Table Grid'
            
            headers = ['', '', 'الرسوم', 'التسجيل', 'المتأخرات', 'المستحق', 'المدفوع', 'المتبقي']
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row_cells = table.add_row().cells
            row_cells[0].text = "المجموع الكلي"
            row_cells[1].text = ""
            for i in range(2, 8):
                value = grand_total[i-2] if grand_total[i-2] is not None else 0
                row_cells[i].text = f"{value:,.2f}"

        # حفظ التقرير في مجلد temp
        temp_dir = os.path.join(app.root_path, 'temp')
        if not os.path.exists(temp_dir):
            try:
                os.makedirs(temp_dir)
            except OSError as e:
                app.logger.error(f"Failed to create temp directory: {str(e)}")
                return jsonify({'error': 'Failed to create temp directory'}), 500

        filename = f'قائمة_غير_الدافعين_{nomcnt}_{noman}.docx'
        filepath = os.path.join(temp_dir, filename)
        
        try:
            doc.save(filepath)
        except Exception as e:
            app.logger.error(f"Failed to save report file: {str(e)}")
            return jsonify({'error': 'Failed to save report file'}), 500

        # إرسال الملف للمستخدم
        try:
            return send_file(
                filepath,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        except Exception as e:
            app.logger.error(f"Failed to send file: {str(e)}")
            return jsonify({'error': 'Failed to send report file'}), 500

    except psycopg2.Error as e:
        app.logger.error(f"Database error in nopayment report: {str(e)}")
        return jsonify({'error': 'Database error occurred'}), 500
    except Exception as e:
        app.logger.error(f"Unexpected error in nopayment report: {str(e)}")
        return jsonify({'error': 'An unexpected error occurred'}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()
#-------------------------------الامتحانات --------------
#-----------------------------------------كشف جمالي الرجات 
@app.route('/imprimtnotecof', methods=['GET', 'POST'])
def imprimtnotecof():
    if request.method == 'GET':
        # عرض النموذج إذا كانت الطريقة GET
        conn = psycopg2.connect(**DB_CONFIG)
        cursor = conn.cursor()        
        # جلب السنوات الدراسية
        cursor.execute("SELECT id_ann, annee_sc FROM ann_scolaire ORDER BY id_ann DESC")
        annees = cursor.fetchall()        
        # جلب الصفوف
        cursor.execute("SELECT id_cl, nom_cl FROM classes ORDER BY nom_cl")
        classes = cursor.fetchall()        
        # جلب العمليات
        cursor.execute("SELECT id_opef, date_opef, nom_periode FROM operation ORDER BY id_opef")
        operations = cursor.fetchall()        
        conn.close()
        
        return render_template('print_final_notes_form.html', 
                            annees=annees,
                            classes=classes,
                            operations=operations,
                            annee_info=session.get('annee_info', {}))
    elif request.method == 'POST':
        try:
            # الحصول على البيانات من النموذج           
            id_opef = session['selected_opef'] 
            nomperiode = request.form.get('nom_periode', session.get('selected_opef_name'))           
            id_cl = session['selected_cl']
            idann = session['annee_info']['id_ann']
            nomclex = request.form.get('nomclex_Entry', '')
            anneesc = request.form.get('anneesc_Entry', '')
            nomcnt = request.form.get('nomcnt_Entry', '')            
            # الاتصال بقاعدة البيانات
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()            
            # جلب المواد
            cursor.execute(
                "SELECT DISTINCT nom_ma FROM matieres NATURAL JOIN classes NATURAL JOIN matiereclasse WHERE id_cl=%s",
                (id_cl,)
            )
            matieres = [row[0] for row in cursor.fetchall()]            
            # جلب الطلاب مع المعدل والترتيب
            cursor.execute("""
                SELECT nom_el, som_notes_period, ranking FROM eleves NATURAL JOIN claselves NATURAL JOIN eleve_note
                WHERE id_cl=%s AND id_ann=%s ORDER BY ranking
            """, (id_cl, idann))
            students_data = cursor.fetchall()
            students = [row[0] for row in students_data]
            student_averages = {row[0]: {'average': row[1], 'ranking': row[2]} for row in students_data}            
            # جلب الدرجات
            cursor.execute("""
                SELECT nom_el, nom_ma, nolesel, note_compo 
                FROM noteesfr NATURAL JOIN matieres NATURAL JOIN classes NATURAL JOIN matiereclasse NATURAL JOIN eleves 
                WHERE id_cl=%s AND id_ope=%s
            """, (id_cl, id_opef))
            schedules = cursor.fetchall()            
            # تنظيم الدرجات
            student_grades = {
                student: {
                    matiere: {'nolesel': '0', 'note_compo': '0'} 
                    for matiere in matieres
                } 
                for student in students
            }            
            for row in schedules:
                if len(row) >= 4:
                    nom_el, nom_ma, nolesel, note_compo = row[0], row[1], row[2], row[3]
                    if nom_el in student_grades and nom_ma in student_grades[nom_el]:
                        student_grades[nom_el][nom_ma]['nolesel'] = str(nolesel) if nolesel is not None else '0'
                        student_grades[nom_el][nom_ma]['note_compo'] = str(note_compo) if note_compo is not None else '0'            
            # إنشاء مستند Word
            doc = Document()            
            # ضبط الصفحة
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)            
            # العناوين
            title = doc.add_paragraph()
            title_run = title.add_run(f"المدرسة: {nomcnt}\t\tالعام الدراسي: {anneesc}\t\tالتاريخ: {datetime.now().strftime('%Y-%m-%d')}")
            title_run.font.name = 'Arial'
            title_run.font.size = Pt(12)
            title_run.bold = True            
            doc.add_heading('كشف الدرجات النهائي', level=1)           
            # إنشاء الجدول
            num_cols = 1 + len(matieres)*2 + 2  # أسماء + مواد + معدل + ترتيب
            num_rows = len(students) + 2  # عناوين + طلاب
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False          
            # العناوين الرئيسية
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'اسم الطالب'            
            # عناوين المواد
            for i, matiere in enumerate(matieres):
                main_cell = hdr_cells[1+i*2]
                main_cell.text = matiere
                main_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                main_cell.merge(hdr_cells[2+i*2])
            # عناوين المعدل والترتيب
            hdr_cells[-2].text = 'المعدل'
            hdr_cells[-1].text = 'الترتيب'

            # العناوين الفرعية
            sub_cells = table.rows[1].cells
            sub_cells[0].text = '---'

            for i in range(len(matieres)):
                sub_cells[1+i*2].text = 'فصل'
                sub_cells[2+i*2].text = 'امتحان'

            sub_cells[-2].text = '---'
            sub_cells[-1].text = '---'
            # تعبئة البيانات
            for row_idx, student in enumerate(students):
                row_cells = table.rows[row_idx+2].cells
                row_cells[0].text = student
                # درجات المواد
                for col_idx, matiere in enumerate(matieres):
                    grades = student_grades[student][matiere]
                    row_cells[1+col_idx*2].text = grades['nolesel']
                    row_cells[2+col_idx*2].text = grades['note_compo']
                # المعدل والترتيب
                row_cells[-2].text = f"{student_averages[student]['average']:.2f}"
                row_cells[-1].text = str(student_averages[student]['ranking'])
                # محاذاة النص
                for cell in row_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # تنسيق الجدول
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
                            if row in [table.rows[0], table.rows[1]]:
                                run.font.bold = True
                                run.font.size = Pt(11)
            # حفظ الملف
            output_dir = os.path.join('Imprimer_Notes')            
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            output_path = os.path.join(output_dir, 'كشف_الدرجات_النهائي.docx')
            doc.save(output_path)            
            # إرجاع الملف للتحميل
            #return send_file(output_path, as_attachment=True)
            flash('تم حفظ كشف الدرجات بنجاح في: ' + output_path, 'success')
            #return redirect(url_for('imprimtnotecof'))
        except Exception as e:
            flash(f'حدث خطأ أثناء إنشاء الملف: {str(e)}', 'danger')
            #return redirect(url_for('imprimtnotecof'))        
        finally:
            if 'conn' in locals():
                conn.close()
#---------------------------------كشف الدرجات بمعدل
#---b-------طباعة كشف درجات الطالب مع imprimnielv + add_note.html----------------------
@app.route('/imprimnielv', methods=['POST'])
def imprimnielv():
    try:
        data = request.json
        student_id = data.get('student_id')
        #id_ope = data.get('id_ope')
        id_ann = data.get('id_ann')
        id_opef = session['selected_opef']
        nomperiode = request.form.get('nom_periode', session.get('selected_opef_name'))
        id_cl = session['selected_cl']
        

        if not student_id or not id_ann:
            return jsonify({"status": "error", "message": "معرّف الطالب أو السنة الدراسية مفقود"}), 400
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        # الحصول على معلومات المركز من قاعدة البيانات
        cur.execute("""
            SELECT id_ann, num_ann, annee_sc, nom_cnt, nom_cntfr, adressear, adressefr, 
            tel, nom_cntfr, nom_cntdet, nom_cntdetfr 
            FROM ann_scolaire NATURAL JOIN nom_centere 
            WHERE id_ann = %s 
        """, (id_ann,))
        ecol_info = cur.fetchone()
        if not ecol_info:
            return jsonify({"status": "error", "message": "لم يتم العثور على معلومات المركز"}), 404
        # استخراج بيانات المركز
        nomcnt = ecol_info[3]
        nomcntfr = ecol_info[4]
        adressear = ecol_info[5]
        adressefr = ecol_info[6]
        tel = ecol_info[7]
        nom_cntfr = ecol_info[8]
        nom_cntdet = ecol_info[9]
        nom_cntdetfr = ecol_info[10]
        # 1. الحصول على بيانات الطالب الأساسية
        query = """
                SELECT nom_ma, nomma_frans, coef_ma, nolesel, note_compo, 
                       round(moyen, 2), round(moy_coeff, 2), etat_ma 
                FROM noteesfr 
                NATURAL JOIN eleves 
                NATURAL JOIN matieres 
                NATURAL JOIN classes 
                NATURAL JOIN claselves 
                WHERE id_el=%s AND id_ope=%s AND id_cl=%s
            """
        cur.execute(query, (student_id, id_opef, id_cl))
        noteelev_list = cur.fetchall()
        # حساب المجاميع
        subtotal0 = sum(item[2] for item in noteelev_list)
        subtotal1 = sum(item[3] for item in noteelev_list)
        subtotal2 = sum(item[4] for item in noteelev_list)
        subtotal3 = sum(item[5] for item in noteelev_list)
        subtotal4 = sum(item[6] for item in noteelev_list)

            # استعلام المعدل والترتيب
        query2 = """
            SELECT round(som_notes_period,2), etat_notes_period, ranking 
            FROM eleve_note 
            NATURAL JOIN eleves 
            NATURAL JOIN classes 
            NATURAL JOIN claselves 
            WHERE id_el=%s AND id_ope=%s AND id_cl=%s
        """
        cur.execute(query2, (student_id, id_opef, id_cl))
        noteelev_list2 = cur.fetchall()

        cur.execute("""
            SELECT id_el, nom_el, nom_cl, annee_sc 
            FROM eleves 
            NATURAL JOIN claselves 
            NATURAL JOIN classes 
            NATURAL JOIN ann_scolaire 
            WHERE id_el = %s AND id_ann = %s
        """, (student_id, id_ann))
        student_info = cur.fetchone()
      
        if not student_info:
            return jsonify({"status": "error", "message": "لم يتم العثور على الطالب"}), 404
        id_el, nomelnote, nomclex, anneesc = student_info
        # استعلام عدد الطلاب
        query3 = """
            SELECT count(id_el) 
            FROM eleves 
            NATURAL JOIN classes 
            NATURAL JOIN claselves 
            WHERE id_cl=%s AND id_ann=%s
        """
        cur.execute(query3, (id_cl, id_ann))
        noteelev_list3 = cur.fetchall()
        # 4. حساب المجاميع
  
        context = {            
            'nomcnt':nomcnt,
            'nomcntfr': nomcntfr,
            'adressear': adressear,
            'adressefr': adressefr,
            'tel': tel,
            'nom_cntfr': nom_cntfr,
            "nom_cntdet": nom_cntdet,
            "nom_cntdetfr": nom_cntdetfr,
            "anneesc": anneesc,
            "nomelnote": nomelnote,
            "nomclex": nomclex,
            "nomperiode": nomperiode,
            "subtotal0": subtotal0,
            "subtotal1": subtotal1,
            "subtotal2": subtotal2,
            "subtotal3": subtotal3,
            "subtotal4": subtotal4,
            "noteelev_list": noteelev_list,
            "noteelev_list2": noteelev_list2,
            "noteelev_list3": noteelev_list3,
        }
        BASE_DIR = Path(__file__).parent
        template_path = BASE_DIR /"templete_Recu"/"note_eleves.docx"
        output_dir = BASE_DIR / "Imprimer_Notes"        
        os.makedirs(output_dir, exist_ok=True)        
        if not template_path.exists():
            return jsonify({
                "status": "error",
                "message": f"ملف القالب غير موجود في: {template_path}"
            }), 404

        # إنشاء وحفظ الوثيقة
        doc = DocxTemplate(str(template_path))
        doc.render(context)
        output_filename = f"كشف الدرجات_.docx"
        output_path = output_dir / output_filename
        doc.save(str(output_path))        
        # إرجاع استجابة JSON فقط
        return jsonify({
            "status": "success",
            
            "message": "تم حفظ الملف بنجاح على الخادم",
            "file_info": {
                "name": output_filename,
                "path": str(output_path),
                "student_name": nomelnote,
                "class": nomclex,
                "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
        }), 200

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        if 'conn' in locals():
            conn.close() 
#---------------------اجمالي كشوفات الطلاب---------------------
@app.route('/imprimclelv', methods=['GET', 'POST'])
def imprimclelv():
    try:
        if request.method == 'GET':
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()
            
            cursor.execute("SELECT id_ann, annee_sc FROM ann_scolaire ORDER BY id_ann DESC")
            annees = cursor.fetchall()
            
            cursor.execute("SELECT id_cl, nom_cl FROM classes ORDER BY nom_cl")
            classes = cursor.fetchall()
            
            cursor.execute("SELECT id_opef, date_opef, nom_periode FROM operation ORDER BY id_opef")
            operations = cursor.fetchall()
            
            conn.close()
            
            return render_template('print_student_notes_form.html', 
                               annees=annees,
                               classes=classes,
                               operations=operations,                               
                               annee_info=session.get('annee_info'))

        elif request.method == 'POST':
            # استقبال البيانات من النموذج
            id_ann = request.form.get('id_ann', session.get('selected_annee', 151))
            id_opef = request.form.get('id_opef', session.get('selected_opef'))
            nomperiode = request.form.get('nom_periode', session.get('selected_opef_name'))
            id_cl = request.form.get('id_cl', session.get('selected_cl'))
            #nomcnt = request.form.get('nomcnt_Entry', '')
            if not all([id_ann, id_opef, id_cl]):
                flash("بيانات ناقصة: يرجى توفير السنة والعملية والفصل", 'error')
                return redirect(url_for('imprimclelv'))

            try:
                con = psycopg2.connect(**DB_CONFIG)
                cur = con.cursor()

                # إنشاء المجلدات إذا لم تكن موجودة
                os.makedirs('templete_Recu', exist_ok=True)
                os.makedirs('Imprimer_Notes', exist_ok=True)

                # جلب بيانات المركز التعليمي
                cur.execute("""
                    SELECT nom_cnt, nom_cntfr, adressear, adressefr, tel, nom_cntdet, nom_cntdetfr 
                    FROM nom_centere 
                    
                """)
                ecol_info = cur.fetchone()

                if not ecol_info:
                    flash("لم يتم العثور على معلومات المركز", 'error')
                    return redirect(url_for('imprimclelv'))
                # جلب معلومات الفصل
                cur.execute("""
                    SELECT nom_cl, annee_sc 
                    FROM classes 
                    NATURAL JOIN ann_scolaire 
                    WHERE id_ann = %s AND id_cl = %s
                """, (id_ann, id_cl))
                class_info = cur.fetchone()
                nomclex, anneesc = class_info if class_info else ("", "")
                # جلب قائمة جميع الطلاب
                cur.execute("""
                    SELECT id_el, nom_el 
                    FROM eleves 
                    NATURAL JOIN claselves 
                    WHERE id_cl = %s AND id_ann = %s
                """, (id_cl, id_ann))
                students = cur.fetchall()

                if not students:
                    flash("لا يوجد طلاب في هذا الفصل", 'error')
                    return redirect(url_for('imprimclelv'))

                template_path = os.path.join('templete_Recu', 'note_eleves2.docx')
                if not os.path.exists(template_path):
                    flash("قالب الوورد غير موجود", 'error')
                    return redirect(url_for('imprimclelv'))
                
                saved_files_count = 0
                
                # جلب عدد الطلاب
                cur.execute("""
                    SELECT count(id_el) FROM eleves 
                    NATURAL JOIN claselves 
                    WHERE id_cl = %s AND id_ann = %s
                """, (id_cl, id_ann))
                total_students = cur.fetchone()[0]
                
                for student in students:
                    id_el, nomelnote = student
                    
                    # جلب درجات الطالب
                    cur.execute("""
                        SELECT nom_ma, nomma_frans, coef_ma, nolesel, note_compo, 
                        ROUND(moyen, 2), ROUND(moy_coeff, 2), etat_ma 
                        FROM noteesfr 
                        NATURAL JOIN eleves 
                        NATURAL JOIN matieres 
                        NATURAL JOIN classes 
                        NATURAL JOIN claselves 
                        WHERE id_ope = %s AND id_cl = %s AND id_el = %s
                    """, (id_opef, id_cl, id_el))
                    noteelev_list = cur.fetchall()

                    # حساب المجاميع
                    subtotals = [sum(col) for col in zip(*[
                        (item[2], item[3], item[4], item[5], item[6]) 
                        for item in noteelev_list
                    ])] if noteelev_list else [0, 0, 0, 0, 0]

                    # جلب التقرير النهائي للطالب
                    cur.execute("""
                        SELECT round(som_notes_period,2), etat_notes_period, ranking 
                        FROM eleve_note 
                        NATURAL JOIN eleves 
                        NATURAL JOIN classes 
                        NATURAL JOIN claselves 
                        WHERE id_ope = %s AND id_cl = %s AND id_el = %s
                    """, (id_opef, id_cl, id_el))
                    noteelev_list2 = cur.fetchall()

                    # إعداد البيانات للقالب
                    student_data = {
                        #nom_cnt, nom_cntfr, adressear, adressefr, tel, nom_cntdet, nom_cntdetfr
                        'nomcnt': ecol_info[0],
                        'nomcntfr': ecol_info[1],
                        'adressear': ecol_info[2],
                        'adressefr': ecol_info[3],
                        'tel': ecol_info[4],
                        'nom_cntdet': ecol_info[5],
                        'nom_cntdetfr': ecol_info[6],
                        #'nomcntfr': ecol_info[1],                                                    
                        'anneesc': anneesc,
                        'nomelnote': nomelnote,
                        'nomclex': nomclex,
                        "nomperiode": nomperiode,
                        'noteelev_list': noteelev_list,
                        'noteelev_list2': noteelev_list2,
                        'noteelev_list3': [(total_students,)],
                        'subtotal0': subtotals[0],
                        'subtotal1': subtotals[1],
                        'subtotal2': subtotals[2],
                        'subtotal3': subtotals[3],
                        'subtotal4': subtotals[4],
                    }

                    # إنشاء المستند
                    doc = DocxTemplate(template_path)
                    doc.render(student_data)
                    
                    # حفظ الملف لكل طالب
                    safe_name = secure_filename(nomelnote)
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    doc_name = f"Note_{safe_name}_{timestamp}.docx"
                    doc_path = os.path.join('Imprimer_Notes', doc_name)
                    doc.save(doc_path)
                    saved_files_count += 1

                cur.close()
                con.close()

                if saved_files_count == 0:
                    flash("لم يتم حفظ أي ملفات", 'error')
                    return redirect(url_for('imprimclelv'))

                # إعادة رسالة نجاح بسيطة دون روابط تحميل
                flash(f"تم حفظ {saved_files_count} كشف درجات بنجاح في مجلد Imprimer_Notes", 'success')
                return redirect(url_for('imprimclelv'))

            except Exception as e:
                app.logger.error(f"Error: {str(e)}\n{traceback.format_exc()}")
                flash(f"حدث خطأ أثناء المعالجة: {str(e)}", 'error')
                return redirect(url_for('imprimclelv'))

    except Exception as e:
        app.logger.error(f"Unexpected error: {str(e)}\n{traceback.format_exc()}")
        flash("حدث خطأ غير متوقع", 'error')
        return redirect(url_for('imprimclelv'))

@app.route('/imprimtnote', methods=['GET', 'POST'])
def imprimtnote():
    if request.method == 'GET':
        # جلب البيانات اللازمة لعرض النموذج
        conn = psycopg2.connect(**DB_CONFIG)
        cursor = conn.cursor()
        
        # جلب قائمة الصفوف
        cursor.execute("SELECT id_cl, nom_cl FROM classes")
        classes = cursor.fetchall()
        
        # جلب قائمة العمليات
        cursor.execute("SELECT id_opef, date_opef, nom_periode FROM operation ORDER BY id_opef")
        operations = cursor.fetchall()
        
        conn.close()
        
        return render_template('print_notes.html', 
                            classes=classes,
                            operations=operations,
                            annee_info=session.get('annee_info', {}))

    elif request.method == 'POST':
        try:
            # التحقق من وجود بيانات الجلسة المطلوبة
            required_session_keys = ['selected_ope', 'selected_cl', 'annee_info']
            if not all(key in session for key in required_session_keys):
                flash('بيانات الجلسة غير كاملة، الرجاء إعادة الاختيار', 'danger')
                return redirect(url_for('imprimtnote'))
            
            id_ope = session['selected_ope']
            id_cl = session['selected_cl']
            idann = session['annee_info']['id_ann']
            nomclex = request.form.get('nomclex_Entry', '')
            anneesc = request.form.get('anneesc_Entry', '')
            nomcnt = request.form.get('nomcnt_Entry', '')

            # استعلامات قاعدة البيانات
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()
            # استعلام المواد الدراسية
            cursor.execute("""
                SELECT DISTINCT nom_ma FROM matieres 
                NATURAL JOIN classes 
                NATURAL JOIN matiereclasse 
                WHERE id_cl = %s
            """, (id_cl,))
            matieres = [row[0] for row in cursor.fetchall()]

            # استعلام أسماء الطلاب مع ترتيبهم
            cursor.execute("""
                SELECT id_el, nom_el FROM eleves 
                NATURAL JOIN classes 
                NATURAL JOIN claselves 
                WHERE id_cl = %s AND id_ann = %s 
                ORDER BY id_el
            """, (id_cl, idann))
            eleves = cursor.fetchall()

            # استعلام الدرجات مع جميع البيانات المطلوبة
            cursor.execute("""
                SELECT id_el, nom_el, nom_ma, nolesel FROM noteesfr NATURAL JOIN eleves NATURAL JOIN matieres 
                WHERE id_cl = %s AND id_ope = %s
                ORDER BY id_el, nom_ma
            """, (id_cl, id_ope))
            notes = cursor.fetchall()
            # إنشاء DataFrame باستخدام pandas
            import pandas as pd            
            # إنشاء قاموس لتخزين الدرجات
            notes_dict = {}
            for eleve in eleves:
                notes_dict[eleve[1]] = {matiere: '' for matiere in matieres}
            # تعبئة الدرجات في القاموس
            for note in notes:
                id_el, nom_el, nom_ma, nolesel = note
                if nom_el in notes_dict and nom_ma in notes_dict[nom_el]:
                    notes_dict[nom_el][nom_ma] = nolesel

            # تحويل القاموس إلى DataFrame
            df = pd.DataFrame.from_dict(notes_dict, orient='index')
            # إضافة عمود لاسم الطالب
            df.reset_index(inplace=True)
            df.rename(columns={'index': 'اسم الطالب'}, inplace=True)
            # إضافة معلومات الرأس كصفوف في الأعلى
            header_info = pd.DataFrame([
                ['', f"المدرسة: {nomcnt}", '', ''],
                ['', f"الصف: {nomclex}", '', ''],
                ['', f"العام الدراسي: {anneesc}", '', '']
            ])
            
            # دمج معلومات الرأس مع البيانات
            final_df = pd.concat([header_info, df], ignore_index=True)
            
            # حفظ الملف Excel
            if not os.path.exists('Imprimer_Notes'):
                os.makedirs('Imprimer_Notes')
            
            file_path = os.path.join('Imprimer_Notes', 'Imprim_Note.xlsx')
            final_df.to_excel(file_path, index=False, header=True)
            
            return send_file(file_path, as_attachment=True)

        except Exception as e:
            flash(f'حدث خطأ أثناء إنشاء الملف: {str(e)}', 'danger')
            return redirect(url_for('imprimtnote'))
        
        finally:
            if 'conn' in locals():
                conn.close()

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))
if __name__ == '__main__':
    app.run(debug=True)

#if __name__ == '__main__':
#    # فتح المتصفح تلقائيًا في خيط منفصل
#    threading.Thread(target=open_browser).start()
#    
#    # تشغيل التطبيق مع دعم HTTPS للتواصل الآمن
#    context = ('cert.pem', 'key.pem')  # شهادات SSL (يمكن إنشاؤها باستخدام OpenSSL)
#    app.run(host='0.0.0.0', port=5000, ssl_context=context, debug=True)
